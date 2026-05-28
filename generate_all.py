"""
Build all.docx from 234_fixed.docx by prepending Abstract+Introduction
and appending Conclusion+References. Uses XML exclusively for consistent order.
"""
import os, copy
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ═══════════════════════════════════════════════════════════════════════════════
# XML helpers
# ═══════════════════════════════════════════════════════════════════════════════
def _make_para(text, style='Normal', bold=False, size=10, align=None, space_after=4, space_before=0):
    """Create a w:p element with run formatting."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style)
    pPr.append(pStyle)
    if space_after or space_before:
        sp = OxmlElement('w:spacing')
        if space_after:
            sp.set(qn('w:after'), str(int(space_after * 20)))
        if space_before:
            sp.set(qn('w:before'), str(int(space_before * 20)))
        sp.set(qn('w:line'), str(int(240 * 1.15)))
        sp.set(qn('w:lineRule'), 'auto')
        pPr.append(sp)
    if align:
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), align)
        pPr.append(jc)
    p.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:ascii'), 'Times New Roman')
    rf.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rf)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(size * 2))
    rPr.append(sz)
    if bold:
        rPr.append(OxmlElement('w:b'))
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    return p

def _make_heading(text, level=1):
    sz = {1: 11, 2: 10, 3: 10}[level]
    return _make_para(text, style=f'Heading {level}', bold=True, size=sz, space_before=8, space_after=4)

def _make_spacer():
    return _make_para('', style='Normal', size=6, space_after=0)

# ═══════════════════════════════════════════════════════════════════════════════
# Load base document
# ═══════════════════════════════════════════════════════════════════════════════
doc = Document(r'E:\DAFB-CLS\papers\234_fixed.docx')
body = doc.element.body
first_el = body[0]  # "II. RELATED WORK" paragraph

# ═══════════════════════════════════════════════════════════════════════════════
# FRONT MATTER: insert Title, Abstract, Introduction before II. RELATED WORK
# Iterate FORWARD so first item ends up earliest in document
# ═══════════════════════════════════════════════════════════════════════════════
front = [
    ('title1', _make_para('DAFB-CLS: Depth-Adaptive Foreground-Background CLS Decoupling',
                          style='Normal', bold=True, size=14, align='center', space_after=0)),
    ('title2', _make_para('for Vision Transformers',
                          style='Normal', bold=True, size=14, align='center', space_after=10)),
    ('sp0', _make_spacer()),
    ('abs_h', _make_heading('Abstract', 1)),
    ('abs1', _make_para(
        'Vision Transformers (ViTs) rely on a learnable [CLS] token to aggregate patch-level '
        'features into a global image representation. However, under standard image-level '
        'supervision, the [CLS] token suffers from Lazy Aggregation: it preferentially attends '
        'to numerous background patches as a computational shortcut, diluting foreground object '
        'signals. Existing post-hoc methods address this by scoring patches with a single '
        'foreground cue and applying a fixed selection budget, but they falter when the cue '
        'is ambiguous or when the fixed budget mismatches object size. Beyond the spatial '
        'dimension, standard residual connections accumulate layer outputs uniformly, ignoring '
        'that foreground and background benefit from representations at different depths.'
    )),
    ('abs2', _make_para(
        'We propose DAFB-CLS, a lightweight post-hoc framework that jointly resolves both '
        'dimensions. Spatially, we fuse four complementary foregroundness cues and learn an '
        'image-adaptive soft mask via a learned per-image threshold, replacing brittle single-cue '
        'scoring. In depth, we decouple the [CLS] token into independent foreground and background '
        'streams, each with its own learned depth-wise attention over multi-layer features, and '
        'fuse them through a task-adaptive gate. DAFB-CLS adds only 0.59M trainable parameters '
        'for ViT-S/16 (under 3% of the frozen backbone) and requires no text supervision.'
    )),
    ('abs3', _make_para(
        'On unsupervised object discovery, DAFB-CLS achieves 79.43% CorLoc and 31.27% Mask IoU '
        'on VOC 2012 (+13.7 pp and +18.0 pp over LaSt-ViT), with consistent cross-dataset gains '
        'on COCO 2017 (+11.4 pp CorLoc, +15.8 pp Mask IoU). On open-vocabulary segmentation, it '
        'reaches 61.70% mIoU and 79.64% PiB. Ablation confirms multi-cue foregroundness drives '
        'spatial localization, adaptive budget prevents mask collapse, and independent depth '
        'attention benefits each stream. Stress tests identify stable background scenes as the '
        'primary failure mode, motivating future work on high-frequency rejection and learned '
        'background priors.'
    )),
    ('sp1', _make_spacer()),
    ('intro_h', _make_heading('I. INTRODUCTION', 1)),
    ('intro1', _make_para(
        'Dense visual understanding tasks — object discovery, weakly-supervised segmentation, '
        'and open-vocabulary localization — require models to identify foreground regions without '
        'pixel-level supervision. Vision Transformers (ViTs) [1] are the dominant backbone for '
        'these tasks, producing a [CLS] token that aggregates patch-level information via '
        'self-attention across all layers.'
    )),
    ('intro2', _make_para(
        'Despite their success, ViTs exhibit a systematic bias: under image-level supervision, '
        'the [CLS] token preferentially attends to background patches — more numerous and '
        'statistically easier to model — as a computational shortcut. This Lazy Aggregation '
        'dilutes the foreground signal needed for dense prediction. Correcting the aggregation '
        'mechanism itself, rather than post-processing its output, provides a more principled '
        'solution.'
    )),
    ('intro3', _make_para(
        'Two lines of work have addressed this. First, LaSt-ViT [2] scores patches by frequency-'
        'domain stability and selects a fixed Top-K for CLS computation. This approach has three '
        'limitations: (i) frequency stability as the sole cue conflates smooth foregrounds with '
        'smooth backgrounds; (ii) fixed Top-K cannot adapt to varying object sizes; (iii) '
        'background information is discarded entirely. Second, Attention Residuals (AttnRes) [4] '
        'learns content-dependent depth attention in NLP, but applies a single pattern to all '
        'tokens. In vision, optimal depth profiles differ for foreground and background. Post-hoc '
        'methods — CAM [5], DINO-seg [6], TokenCut [7] — inherit the biased aggregation.'
    )),
    ('intro4', _make_para(
        'We propose DAFB-CLS (Depth-Adaptive Foreground-Background CLS Decoupling), a '
        'lightweight post-hoc framework that jointly resolves both spatial and depth-dimensional '
        'limitations (Fig. 1). The key innovation is twofold. First, for the spatial dimension, '
        'we fuse four complementary foregroundness cues — frequency stability, depth consistency, '
        'semantic alignment, and spatial compactness — into an image-adaptive soft mask, replacing '
        'brittle single-cue scoring. Second, for the depth dimension, we decouple the [CLS] token '
        'into independent foreground and background streams with separate learned depth-wise '
        'attention, and fuse them through a task-adaptive gate.'
    )),
    ('intro5', _make_para(
        'On VOC 2012, DAFB-CLS achieves 79.43% CorLoc and 31.27% Mask IoU (+13.7 pp and +18.0 pp '
        'over LaSt-ViT), with consistent cross-dataset generalization to COCO 2017. Our '
        'contributions are:'
    )),
    ('c1', _make_para(
        '• A multi-cue foregroundness framework with image-adaptive soft masking, '
        'substantially improving spatial foreground precision over single-cue alternatives.'
    )),
    ('c2', _make_para(
        '• A dual CLS decoupling mechanism with independent depth-wise attention for '
        'foreground and background streams.'
    )),
    ('c3', _make_para(
        '• Comprehensive experiments over six baselines on two benchmarks, with ablation '
        'studies and stress tests characterizing each component and failure mode.'
    )),
    ('c4', _make_para(
        '• An efficient post-hoc design (0.59M trainable parameters, under 3% of backbone) '
        'practical as a drop-in enhancement for pretrained ViTs.'
    )),
    ('sp2', _make_spacer()),
]

for key, el in front:
    first_el.addprevious(el)
print('Prepended Title, Abstract, Introduction')

# ═══════════════════════════════════════════════════════════════════════════════
# BACK MATTER: append Conclusion + References to end of body
# ═══════════════════════════════════════════════════════════════════════════════
back = [
    ('conc_h', _make_heading('V. CONCLUSION', 1)),
    ('conc1', _make_para(
        'This paper addressed Lazy Aggregation in Vision Transformers, where the [CLS] token '
        'systematically biases toward background patches. We proposed DAFB-CLS, a lightweight '
        'post-hoc framework that decouples the [CLS] token into independent foreground and '
        'background streams with depth-adaptive aggregation, and enriches the foreground signal '
        'through multi-cue scoring with image-adaptive soft masking.'
    )),
    ('conc2', _make_para(
        'The key insight is that foreground-background separation must occur at the aggregation '
        'level — accounting for both spatial selection (which patches are foreground) and depth '
        'selection (which layers best represent each stream). By fusing four complementary cues, '
        'learning per-image thresholds, and maintaining independent depth attention, DAFB-CLS '
        'achieves +13.7 pp CorLoc and +18.0 pp Mask IoU over the strongest aggregation baseline, '
        'with robust cross-dataset generalization. The ablation reveals a clear division of labor '
        'among components, while stress tests identify stable backgrounds as the primary remaining '
        'challenge. This architectural limitation points toward explicit high-frequency rejection '
        'and learned background priors as promising future directions.'
    )),
    ('sp3', _make_spacer()),
    ('ref_h', _make_heading('References', 1)),
]

ref_entries = [
    '[1] A. Dosovitsky, L. Beyer, A. Kolesnikov, et al., "An Image is Worth 16x16 Words: Transformers for Image Recognition at Scale," in Proc. ICLR, 2021.',
    '[2] S. Shi et al., "LaSt-ViT: Rethinking CLS Token Aggregation in Vision Transformers," in Proc. ECCV, 2024.',
    '[3] T. He et al., "On the Dilution of Residual Connections in Pre-Norm Transformers," in Proc. ICLR, 2025.',
    '[4] Kimi Team, "Attention Residuals: Learning Content-Dependent Depth Attention," arXiv:2603.15031, 2026.',
    '[5] B. Zhou et al., "Learning Deep Features for Discriminative Localization," in Proc. CVPR, 2016.',
    '[6] M. Caron et al., "Emerging Properties in Self-Supervised Vision Transformers," in Proc. ICCV, 2021.',
    '[7] Y. Wang et al., "TokenCut: Segmenting Objects in Images and Videos with Self-Supervised Transformer and Normalized Cut," IEEE Trans. Pattern Anal. Mach. Intell., 2023.',
    '[8] A. Radford et al., "Learning Transferable Visual Models from Natural Language Supervision," in Proc. ICML, 2021.',
    '[9] C. Schuhmann et al., "LAION-5B: An Open Large-Scale Dataset for Training Next Generation Image-Text Models," in Proc. NeurIPS, 2022.',
    '[10] R. R. Selvaraju et al., "Grad-CAM: Visual Explanations from Deep Networks via Gradient-Based Localization," in Proc. ICCV, 2017.',
    '[11] M. Everingham et al., "The PASCAL Visual Object Classes (VOC) Challenge," Int. J. Comput. Vis., 2010.',
    '[12] T.-Y. Lin et al., "Microsoft COCO: Common Objects in Context," in Proc. ECCV, 2014.',
]

for i, ref_text in enumerate(ref_entries):
    back.append((f'ref{i}', _make_para(ref_text, size=8.5, space_after=2)))

for key, el in back:
    body.append(el)
print('Appended Conclusion + References')

# ═══════════════════════════════════════════════════════════════════════════════
# Save
# ═══════════════════════════════════════════════════════════════════════════════
out = r'E:\DAFB-CLS\papers\all.docx'
doc.save(out)

# Verify
v = Document(out)
print(f'\nSaved: {out}')
print(f'Paragraphs: {len(v.paragraphs)}, Tables: {len(v.tables)}, Images: {len(v.inline_shapes)}')
for i, p in enumerate(v.paragraphs):
    t = p.text.strip()
    if t in ['Abstract', 'I. INTRODUCTION', 'II. RELATED WORK', 'III. METHOD',
             'IV. EXPERIMENTS', 'V. CONCLUSION', 'References']:
        print(f'  [{i}] {t}')
    elif 'DAFB-CLS: Depth' in t:
        print(f'  [{i}] TITLE')
