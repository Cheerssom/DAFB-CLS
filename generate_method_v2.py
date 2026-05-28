"""
重写 Method 章节 —— 更清晰、更易懂
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Microsoft YaHei'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    hs.font.color.rgb = RGBColor(0, 51, 102)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
    return table

def bp(doc, text):
    doc.add_paragraph(text, style='List Bullet')

def add_formula(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(80, 80, 80)
    return p

def cite(n, name=""):
    if name:
        return f"[{n}, {name}]"
    return f"[{n}]"

# 引用常量
C_VIT = cite(1, 'ViT')
C_DINO = cite(3, 'DINO')
C_DINOV2 = cite(4, 'DINOv2')
C_CLIP = cite(5, 'CLIP')
C_OPENCLIP = cite(6, 'OpenCLIP')
C_REGISTER = cite(15, 'Registers')
C_REGISTERS = cite(16, 'MoreThanRegisters')
C_ATTNRES = cite(17, 'AttnRes')
C_LASTVIT = cite(22, 'LaSt-ViT')

# ==================== METHOD ====================
title = doc.add_heading('Method', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# --- 3.1 问题定义 ---
doc.add_heading('3.1 问题定义与动机', level=1)

doc.add_paragraph(
    f'在标准Vision Transformer {C_VIT}中，输入图像被分割为N个patch（默认14x14=196个），'
    f'每个patch被编码为D维向量。此外，一个特殊的CLS token被添加到序列开头，'
    f'通过L层Transformer的自注意力机制与所有patch交互，最终汇聚为一个D维的全局表示向量。'
    f'这个CLS token被用于下游任务（分类、分割等）。'
)

doc.add_paragraph(
    f'然而，CLS token存在两个核心问题：'
)

doc.add_paragraph(
    f'第一，惰性聚合{C_REGISTER}：CLS token在聚合信息时，会不加区分地吸收所有patch的信息，'
    f'包括大量的背景patch。这导致CLS token的表示中混入了大量背景噪声，'
    f'降低了对前景目标的聚焦能力。Register方法{C_REGISTERS}试图通过添加额外的token来缓解这一问题，'
    f'但需要重新训练模型。'
)

doc.add_paragraph(
    f'第二，惰性累加{C_ATTNRES}：ViT的每一层都会产生一个CLS token的中间表示，'
    f'最终的CLS token是所有层中间表示的等权重叠加。这意味着模型无法自适应地决定'
    f'哪些层的输出对当前任务更重要——浅层的空间细节和深层的语义信息被同等对待。'
)

doc.add_paragraph(
    f'DAFB-CLS的核心思想是：既然CLS token混合了前景和背景信息，'
    f'我们就把它拆成两个独立的token——一个专门编码前景信息，一个专门编码背景信息；'
    f'既然各层的贡献被等权累加，我们就用注意力机制让模型自适应地选择每层的重要性。'
    f'整个过程不需要修改或重新训练backbone，只需要在冻结的backbone之上训练轻量的聚合器模块。'
)

# --- 3.2 整体流程 ---
doc.add_heading('3.2 整体流程', level=1)

doc.add_paragraph(
    f'DAFB-CLS的前向传播分为五个阶段，可以用一个简单的比喻来理解：'
)

doc.add_paragraph(
    f'想象你站在一幅画前面，需要判断画中哪些部分是"主体"（前景），哪些是"背景"。'
)

bp(doc, f'阶段一（多线索前景区分）：你会同时使用多种线索来判断——'
   f'画面中哪些区域有丰富的纹理和边缘（频率线索），'
   f'哪些区域在不同角度看都保持一致（深度一致性线索），'
   f'哪些区域与整幅画的主题最相关（语义对齐线索），'
   f'哪些区域在空间上连成一片（空间紧凑性线索）。'
   f'你不会只依赖单一线索，而是综合所有线索做出判断。')

bp(doc, f'阶段二（自适应前景预算）：你会根据每幅画的具体情况，'
   f'自适应地决定"前景大概占画面的多大面积"。'
   f'一幅画可能前景占30%，另一幅可能占60%——不能用一个固定的比例来套所有画。')

bp(doc, f'阶段三（双CLS解耦）：确定了前景和背景区域后，'
   f'你分别把前景区域的信息汇总成一个"前景摘要"，'
   f'把背景区域的信息汇总成一个"背景摘要"。'
   f'这两个摘要各自独立，互不干扰。')

bp(doc, f'阶段四（深度注意力）：你在不同距离（相当于ViT的不同层）观察这幅画，'
   f'每个距离看到的细节不同。前景摘要可能从近距离观察更有用，'
   f'背景摘要可能从远距离观察更有用。你用注意力机制自适应地选择每个摘要的最佳观察距离。')

bp(doc, f'阶段五（任务自适应融合）：最后，你根据具体任务（是要分类还是要分割），'
   f'自适应地决定最终结论应该更依赖前景摘要还是背景摘要。')

doc.add_paragraph(
    f'形式化地，给定输入图像I，整个流程可以用以下公式概括：'
)

add_formula(doc, f'阶段零（特征提取）：{{x^l}}_{{l=1}}^L, c = Backbone(I)')
add_formula(doc, f'阶段一（前景区分）：F = CueFusion(x^L)  →  每个patch的前景性分数 [B, N]')
add_formula(doc, f'阶段二（自适应掩码）：m^F, m^B = MaskPrediction(F, c)  →  前景/背景掩码 [B, N]')
add_formula(doc, f'阶段三（双CLS聚合）：B_l^F = WeightedAvg(x^l, m^F),  B_l^B = WeightedAvg(x^l, m^B)')
add_formula(doc, f'阶段四（深度注意力）：C_F = DepthAttn({{B_l^F}}),  C_B = DepthAttn({{B_l^B}})')
add_formula(doc, f'阶段五（融合输出）：C = g * C_F + (1-g) * C_B')

doc.add_paragraph(
    f'其中x^l是第l层的patch特征（形状[B, N, D]，B是batch大小，N是patch数，D是特征维度），'
    f'c是全局CLS token（形状[B, D]），L是提取的层数（默认提取第3、6、9、12层，共4层）。'
    f'下面逐一详细介绍每个阶段。'
)

# --- 3.3 阶段一 ---
doc.add_heading('3.3 阶段一：多线索前景区分度计算', level=1)

doc.add_paragraph(
    f'这一步的目标是：为图像中的每个patch计算一个"前景性分数"，'
    f'分数越高表示该patch越可能属于前景目标。'
)

doc.add_paragraph(
    f'现有方法的不足：LaSt-ViT {C_LASTVIT}仅使用频率稳定性一个信号，'
    f'但当背景是平滑的（如天空、白色墙壁）时，背景patch的频率稳定性也很高，'
    f'会被误判为前景。单一信号无法可靠地区分前景和背景。'
)

doc.add_paragraph(
    f'我们的解决方案：同时使用四种互补的线索，每种基于不同的物理直觉。'
)

# 线索1
doc.add_paragraph()
r = doc.add_paragraph().add_run('线索一：频率稳定性（FrequencyStabilityCue）')
r.bold = True

doc.add_paragraph(
    f'直觉：前景目标（如动物、汽车）的特征通常包含丰富的高频成分（边缘、纹理），'
    f'而背景区域（如天空、水面）的特征更平滑，能量集中在低频。'
)

doc.add_paragraph('具体做法：对每个patch的D维特征做FFT变换，用高斯低通滤波器滤掉高频，'
   '然后比较原始特征和滤波后特征的差异：')

add_formula(doc, 'x_fft = FFT(x_i^l)                              -- 频域变换')
add_formula(doc, 'H(f) = exp(-0.5 * (f / 0.25)^2)                 -- 高斯低通滤波器')
add_formula(doc, 'x_filtered = IFFT(x_fft * H(f)).real             -- 滤波后特征')
add_formula(doc, 's_freq = mean( x_filtered / |x_filtered - x| )   -- 稳定性分数')

doc.add_paragraph(
    '当patch特征平滑时（背景），x_filtered约等于x，分母很小，分数高；'
    '当patch特征包含高频成分时（前景边缘），|x_filtered - x|大，分数低。'
    '所以这个分数的含义是：越平滑的patch分数越高，越可能是背景。'
    '我们取最后一层的分数作为该线索的输出。'
)

# 线索2
doc.add_paragraph()
r = doc.add_paragraph().add_run('线索二：深度一致性（DepthConsistencyCue）')
r.bold = True

doc.add_paragraph(
    '直觉：前景目标（如一只猫）在ViT的不同层都会被持续关注，'
    '其patch特征在不同层之间保持较高的相似性；'
    '而背景区域在不同层的特征可能变化较大（浅层是纹理，深层变成均匀的背景语义）。'
)

doc.add_paragraph('具体做法：计算每个patch在所有层的特征与跨层平均特征的余弦相似度：')

add_formula(doc, 'x_mean = mean over L layers of x_i^l        -- 跨层平均特征')
add_formula(doc, 's_depth = mean over L of cosine(x_i^l, x_mean)  -- 一致性分数')

doc.add_paragraph(
    '前景patch的跨层一致性高，得分高；背景patch的跨层一致性低，得分低。'
)

# 线索3
doc.add_paragraph()
r = doc.add_paragraph().add_run('线索三：语义对齐（SemanticAlignmentCue）')
r.bold = True

doc.add_paragraph(
    f'直觉：CLS token编码了整幅图像的语义信息，'
    f'与CLS token相似度高的patch更可能与图像主题（通常是前景目标）相关。'
    f'对于OpenCLIP {C_OPENCLIP}骨干，还可以利用文本嵌入——'
    f'计算每个patch与各类别文本描述（如"a photo of a cat"）的相似度。'
)

add_formula(doc, 's_sem = mean over L of cosine(x_i^l, c)     -- patch与CLS的相似度')

# 线索4
doc.add_paragraph()
r = doc.add_paragraph().add_run('线索四：空间紧凑性（SpatialCompactnessCue）')
r.bold = True

doc.add_paragraph(
    '直觉：前景目标在图像中通常是空间连续的（一只猫占据一片连续区域），'
    '而孤立的高分patch更可能是噪声。'
    '因此，对语义对齐分数做一次空间平滑，抑制孤立的高分patch，鼓励空间连续的前景区域。'
)

add_formula(doc, 's_spatial = AvgPool2d( reshape(s_sem, 14, 14), kernel_size=3 )')

# 融合
doc.add_paragraph()
r = doc.add_paragraph().add_run('四条线索的融合方式')
r.bold = True

doc.add_paragraph(
    '四条线索各自提供了不同角度的前景判断。我们通过两种互补的方式将它们融合：'
)

doc.add_paragraph(
    '方式一（线性融合）：为每条线索分配一个可学习的权重，权重通过softmax归一化后加权求和。'
    '初始时四条线索权重相等（各0.25），训练过程中模型会自动学习每条线索的重要性。'
)
add_formula(doc, 'w = softmax(w_1, w_2, w_3, w_4)')
add_formula(doc, 'F_linear = w_1*s_freq + w_2*s_depth + w_3*s_sem + w_4*s_spatial')

doc.add_paragraph(
    '方式二（非线性修正）：将四条线索拼接为4维向量，通过一个3层MLP（4→256→256→1）'
    '学习线索之间的非线性交互关系。例如，模型可以学习到'
    '"当频率稳定性高且语义对齐也高时，该patch更可能是前景"这样的组合规则。'
)
add_formula(doc, 'F_mlp = MLP([s_freq; s_depth; s_sem; s_spatial])')

doc.add_paragraph(
    '两种方式的结果相加，再加上一个空间平滑残差项（3x3卷积 + Sigmoid），'
    '得到最终的前景性分数F：'
)
add_formula(doc, 'F = F_linear + F_mlp + 0.5 * Sigmoid(Conv2d(F, kernel=3))')

doc.add_paragraph(
    '空间平滑残差的作用是：在最终分数上施加一层空间连续性约束，'
    '使得相邻patch的前景性分数趋于一致。权重0.5是固定的，控制平滑的强度。'
)

# --- 3.4 阶段二 ---
doc.add_heading('3.4 阶段二：自适应前景预算（掩码生成）', level=1)

doc.add_paragraph(
    f'上一步得到了每个patch的连续前景性分数F（范围大致在0~1之间），'
    f'但这还不是二值化的前景/背景掩码。我们需要一个阈值来决定"分数高于多少算前景"。'
)

doc.add_paragraph(
    f'现有方法的不足：LaSt-ViT {C_LASTVIT}使用固定的Top-K选择——'
    f'无论图像内容如何，总是选择分数最高的30%的patch作为前景。'
    f'这导致两个问题：（1）当前景实际只占10%时，30%的选择会包含大量背景；'
    f'（2）当前景占50%时，30%的选择会遗漏部分前景。'
)

doc.add_paragraph(
    f'我们的解决方案：用一个小型神经网络从全局CLS特征中预测每张图像的最优阈值τ。'
    f'这样，模型可以根据图像的具体内容自适应地调整阈值。'
)

doc.add_paragraph('阈值预测网络（AdaptiveBudgetModule）：')

add_formula(doc, '输入：全局CLS特征 c ∈ R^D')
add_formula(doc, 'τ = Linear(D, 128) → GELU → Linear(128, 1) → 输出标量阈值')

doc.add_paragraph(
    '可选地，还可以将前景性分数的统计量（均值、标准差、最大值、最小值）'
    '拼接到输入中，为阈值预测提供更多参考信息。'
)

doc.add_paragraph('得到阈值τ后，通过sigmoid函数生成软前景掩码：')
add_formula(doc, 'm^F = sigmoid((F - τ) / T)')

doc.add_paragraph(
    '其中T是可学习的温度参数（初始值0.1），控制sigmoid的陡峭程度。'
    'T越小，掩码越接近二值化（0或1）；T越大，掩码越平滑（中间值越多）。'
    '软掩码的优势在于保持可微性，且在掩码边界产生平滑过渡，避免硬阈值的锯齿伪影。'
)

doc.add_paragraph()
r = doc.add_paragraph().add_run('背景掩码的生成（BackgroundnessHead）：')
r.bold = True

doc.add_paragraph(
    '前景掩码确定后，背景掩码不是简单地取补集（1-m^F），'
    '而是通过独立的神经网络从patch特征中直接学习背景概率：'
)

add_formula(doc, 'bg_score = sigmoid(MLP_bg(x^L))          -- 背景概率')
add_formula(doc, 'uncertainty = sigmoid(MLP_unc(x^L))       -- 不确定性估计')
add_formula(doc, 'm^B = bg_score * (1 - m^F) * uncertainty  -- 背景掩码')

doc.add_paragraph(
    '背景掩码是三个因子的乘积：'
    '（1）bg_score——模型学习的背景概率，高分patch更可能是背景；'
    '（2）(1-m^F)——前景掩码的补集，确保前景和背景不重叠；'
    '（3）uncertainty——不确定性门控，当模型对某patch是否为背景不确定时，'
    '降低该patch在背景聚合中的权重。'
)

# --- 3.5 阶段三 ---
doc.add_heading('3.5 阶段三：双CLS解耦聚合', level=1)

doc.add_paragraph(
    f'这一步是DAFB-CLS的核心创新。'
    f'有了前景掩码m^F和背景掩码m^B后，我们分别计算前景CLS token和背景CLS token。'
)

doc.add_paragraph(
    f'现有方法的不足：Register {C_REGISTER}只是被动地消除artifact，'
    f'没有主动地将前景和背景信息分开编码；'
    f'Attention Residuals {C_ATTNRES}虽然改进了深度聚合，'
    f'但仍是在单一的混合CLS token上操作，前景和背景信息互相干扰。'
)

doc.add_paragraph(
    f'我们的做法：对每个提取的ViT层，分别用前景掩码和背景掩码对patch特征进行加权平均，'
    f'生成该层的前景CLS token和背景CLS token：'
)

add_formula(doc, 'B_l^F = Σ_i (m_i^F * x_i^l) / Σ_i m_i^F    -- 第l层的前景CLS token')
add_formula(doc, 'B_l^B = Σ_i (m_i^B * x_i^l) / Σ_i m_i^B    -- 第l层的背景CLS token')

doc.add_paragraph(
    '直觉理解：假设一张图中有一只猫（前景）和一面墙（背景）。'
    '前景CLS token B_l^F 只包含猫的特征信息（因为只有猫所在patch的掩码值高），'
    '背景CLS token B_l^B 只包含墙的特征信息。'
    '两者互不干扰，各自独立地编码了前景和背景的语义。'
)

doc.add_paragraph(
    '对提取的L=4个层（第3、6、9、12层）分别执行这一操作，'
    '得到前景CLS序列 B^F = {B_3^F, B_6^F, B_9^F, B_12^F} 和'
    '背景CLS序列 B^B = {B_3^B, B_6^B, B_9^B, B_12^B}，'
    '各自形状为[B, L, D]。'
)

# --- 3.6 阶段四 ---
doc.add_heading('3.6 阶段四：深度自适应注意力', level=1)

doc.add_paragraph(
    f'上一步得到了前景和背景各自在L=4个层上的CLS token序列。'
    f'但不同层的CLS token质量不同——浅层（如第3层）包含更多空间细节，'
    f'深层（如第12层）包含更多语义信息。我们需要自适应地选择每条流的最佳层组合。'
)

doc.add_paragraph(
    f'现有方法的不足：Attention Residuals {C_ATTNRES}提出了深度注意力的概念，'
    f'但只在单一的混合CLS token上操作，前景和背景共享同一套深度权重。'
    f'然而，前景信息可能更依赖深层语义（如"这是一只猫"），'
    f'而背景信息可能更依赖浅层纹理（如"这是草地"），两者需要不同的深度权重分布。'
)

doc.add_paragraph(
    f'我们的做法：为前景和背景流各自维护一个独立的注意力模块。'
    f'每个模块包含一个可学习的query向量q和一个RMSNorm归一化层：'
)

add_formula(doc, '前景流：')
add_formula(doc, '  keys_F = RMSNorm(B^F)                    -- 归一化后的前景CLS序列 [B, L, D]')
add_formula(doc, '  scores_F = q_F^T * keys_F                 -- query与各层key的相似度 [B, L]')
add_formula(doc, '  beta_F = softmax(scores_F)                 -- 深度注意力权重 [B, L]')
add_formula(doc, '  C_F = Σ_l (beta_F_l * B_l^F)              -- 前景CLS汇总 [B, D]')

add_formula(doc, '背景流：（完全相同的结构，但参数独立）')
add_formula(doc, '  keys_B = RMSNorm(B^B)')
add_formula(doc, '  scores_B = q_B^T * keys_B')
add_formula(doc, '  beta_B = softmax(scores_B)')
add_formula(doc, '  C_B = Σ_l (beta_B_l * B_l^B)              -- 背景CLS汇总 [B, D]')

doc.add_paragraph(
    '关键设计细节：'
)

bp(doc, '零初始化：query向量q_F和q_B初始化为全零向量。'
   '这意味着初始时所有层的注意力分数相同（=0），softmax后均匀分布（=1/L）。'
   '这是一种安全的起始点——初始时模型等权聚合所有层（等价于标准残差连接），'
   '训练过程中逐渐学习差异化的权重。这避免了随机初始化可能带来的训练不稳定。')

bp(doc, 'RMSNorm：对keys进行Root Mean Square归一化（x / sqrt(mean(x^2)) * weight），'
   '比LayerNorm更简洁稳定，遵循现代LLM的设计实践。'
   '归一化在float32精度下进行以确保数值稳定。')

bp(doc, '独立参数：前景流和背景流的query向量和RMSNorm参数完全独立，互不共享。'
   '这允许它们各自学习最适合的深度权重分布。')

# --- 3.7 阶段五 ---
doc.add_heading('3.7 阶段五：任务自适应融合', level=1)

doc.add_paragraph(
    '经过前四个阶段，我们得到了两个独立的CLS汇总向量：'
    'C_F（前景CLS，只包含前景信息）和 C_B（背景CLS，只包含背景信息）。'
    '最后一步是将它们融合为一个统一的CLS表示，用于下游任务。'
)

doc.add_paragraph(
    '不同任务对前景和背景的需求不同：'
    '物体发现任务主要依赖前景信息，背景信息干扰定位；'
    '分割任务需要同时利用前景和背景信息来区分目标和非目标区域；'
    '分类任务可能需要背景信息作为上下文线索（如"在水面上"暗示"船"）。'
)

doc.add_paragraph(
    '因此，我们设计了一个门控网络，根据图像内容和任务需求自适应地调整融合比例：'
)

add_formula(doc, 'g = sigmoid( MLP( [C_F; C_B; c] ) )     -- 标量门控值，范围(0, 1)')
add_formula(doc, 'C = g * C_F + (1 - g) * C_B              -- 最终CLS表示')

doc.add_paragraph(
    '其中c是backbone的原始全局CLS token，作为额外的上下文信息。'
    'MLP结构为 Linear(3D, 256) → GELU → Linear(256, 1)。'
    '当g趋近1时，最终输出以前景为主；当g趋近0时，以背景为主。'
    '这个门控值是逐图像预测的，不同图像可以有不同的融合比例。'
)

# --- 3.8 损失函数 ---
doc.add_heading('3.8 损失函数', level=1)

doc.add_paragraph(
    'DAFB-CLS的训练损失由四项组成，每项对应一个训练目标：'
)

add_formula(doc, 'L_total = L_task + λ_fg * L_mask + λ_dec * L_decouple + λ_bud * L_budget')

doc.add_paragraph()
r = doc.add_paragraph().add_run('（1）主任务损失 L_task：')
r.bold = True
doc.add_paragraph(
    '与最终任务目标直接相关的损失。分类任务使用交叉熵损失；'
    '分割任务使用patch-text相似度的交叉熵损失；'
    '物体发现任务使用patch得分与GT边界框的交叉熵损失。'
    '这一损失驱动整个框架朝正确的任务目标优化。'
)

doc.add_paragraph()
r = doc.add_paragraph().add_run('（2）前景掩码损失 L_mask：')
r.bold = True
doc.add_paragraph(
    '监督预测的前景掩码m^F与真实前景区域的一致性。'
    '使用BCE损失和Dice损失的组合：'
)
add_formula(doc, 'L_mask = BCE(m^F, m_gt) + (1 - Dice(m^F, m_gt))')
doc.add_paragraph(
    '由于训练数据可能没有逐pixel的前景标注，'
    '我们从patch分数中生成伪GT掩码：选择分数高于"均值+标准差"的patch作为前景。'
    '这一损失确保前景掩码的空间质量。'
)

doc.add_paragraph()
r = doc.add_paragraph().add_run('（3）解耦损失 L_decouple：')
r.bold = True
doc.add_paragraph(
    '推远前景CLS token C_F和背景CLS token C_B的方向，确保它们编码不同的语义信息。'
    '使用余弦平方损失：'
)
add_formula(doc, 'L_decouple = mean( cosine(C_F, C_B)^2 )')
doc.add_paragraph(
    '当C_F和C_B的方向完全不同时，余弦相似度为0，损失为0；'
    '当它们方向相同时，余弦相似度为1，损失为1。'
    '这鼓励前景和背景token编码互补的信息。'
)

doc.add_paragraph()
r = doc.add_paragraph().add_run('（4）预算正则化损失 L_budget：')
r.bold = True
doc.add_paragraph(
    '将前景掩码的比例约束在合理范围内，防止掩码坍缩（全0或全1）。'
)
add_formula(doc, 'r = mean(m^F)                               -- 前景比例')
add_formula(doc, 'L_budget = max(r_min - r, 0)^2 + max(r - r_max, 0)^2')
doc.add_paragraph(
    '其中r_min=0.1，r_max=0.7。当前景比例低于10%或高于70%时才会产生惩罚。'
    '消融实验表明，去除这一损失会导致Mask IoU从31.27%暴跌至13.91%——'
    '没有约束时，模型会把掩码预测为几乎全1或全0。'
)

# --- 3.9 训练策略 ---
doc.add_heading('3.9 训练策略与实现细节', level=1)

doc.add_paragraph(
    f'DAFB-CLS采用后训练策略：backbone（如DINO {C_DINO}或OpenCLIP {C_OPENCLIP}）完全冻结，'
    f'仅训练聚合器模块（前景区分、掩码生成、双CLS聚合、深度注意力、融合门控）。'
    f'这使得训练成本极低（聚合器参数量远小于backbone），'
    f'且可以即插即用到任何已有的预训练ViT模型上。'
)

doc.add_paragraph('训练使用AdamW优化器，余弦退火学习率调度，梯度裁剪（max_norm=1.0）。'
   '对于COCO数据集，由于标注更复杂（80类、多实例），'
   '需要降低掩码损失权重和学习率，并关闭混合精度训练以避免fp16溢出。')

add_table(doc,
    ['超参数', 'DINO ViT-S/16', 'OpenCLIP ViT-B/16'],
    [
        ['提取的层', '[3, 6, 9, 12]', '[3, 6, 9, 12]'],
        ['隐藏维度', '256', '256'],
        ['学习率', '1e-4 (VOC) / 5e-5 (COCO)', '5e-5'],
        ['batch_size', '16', '16'],
        ['训练轮数', '50', '50'],
        ['掩码损失权重 λ_fg', '1.0 (VOC) / 0.05 (COCO)', '2.0'],
        ['解耦损失权重 λ_dec', '0.05', '0.1'],
        ['预算损失权重 λ_bud', '0.01', '0.02'],
        ['前景比例约束', 'r_min=0.1, r_max=0.7', 'r_min=0.1, r_max=0.7'],
        ['温度初始值', '0.1', '0.1'],
        ['频率sigma', '0.25', '0.25'],
    ]
)

# Save
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'DAFB_CLS_Method_v2.docx')
doc.save(output_path)
print(f'已保存: {output_path}')
