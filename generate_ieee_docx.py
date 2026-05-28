# -*- coding: utf-8 -*-
"""
Generate IEEE-formatted Related Work + Methods docx using the Transactions template.
"""

TEMPLATE = r"E:\DAFB-CLS\参考文献\模板\Transactions-template-and-instructions-on-how-to-create-your-article-formatted.docx"
OUTPUT   = r"E:\DAFB-CLS\DAFB_CLS_RelatedWork_Methods_IEEE.docx"

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, re

def get_style(doc, name):
    """Get a style from the template, falling back gracefully."""
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles["Normal"]

def add_styled_para(doc, text, style_name="Text", bold=False, italic=False,
                    size=None, space_after=None, space_before=None,
                    alignment=None, first_indent=None):
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    pf = p.paragraph_format
    if space_after is not None:
        pf.space_after = Pt(space_after)
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if alignment is not None:
        p.alignment = alignment
    if first_indent is not None:
        pf.first_line_indent = Cm(first_indent)
    return p

def clear_doc(doc):
    """Remove all paragraphs except the last (which Word requires)."""
    body = doc.element.body
    for p in body.findall(qn("w:p"))[1:]:
        body.remove(p)
    for t in body.findall(qn("w:tbl")):
        body.remove(t)

def add_heading(doc, text, level=1):
    style_map = {1: "heading 1", 2: "heading 2", 3: "heading 3"}
    style_name = style_map.get(level, "heading 1")
    p = doc.add_paragraph(text, style=style_name)
    return p

def add_text(doc, text, style="Text", bold=False, italic=False, alignment=None, first_indent=0.75):
    return add_styled_para(doc, text, style_name=style, bold=bold, italic=italic,
                           alignment=alignment, first_indent=first_indent)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        table.style = "Normal Table"
    # header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(8)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(8)
            if j > 0:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table

def write_related_work(doc):
    add_heading(doc, "RELATED WORK", 1)

    # ── 2.1 ──
    add_heading(doc, "Vision Transformers and Dense Feature Representations", 2)

    add_heading(doc, "From Global Attention to Structured Variants", 3)
    add_text(doc,
        "The Vision Transformer (ViT) [1] demonstrated that pure attention-based architectures "
        "can match or exceed convolutional networks on image classification when pretrained at "
        "sufficient scale. By treating image patches as tokens and appending a learnable [CLS] token, "
        "ViT leverages global self-attention to build a holistic image representation. This design "
        "has since become the foundation of numerous vision and multimodal models, including CLIP [2], "
        "DINO [3], and BEiT [4].")
    add_text(doc,
        "However, global quadratic attention also introduces substantial computational overhead, "
        "motivating a line of work on structured sparse attention. PVT [5] progressively reduces "
        "spatial resolution across stages via spatial-reduction attention, while Swin Transformer [6] "
        "restricts attention within local windows with periodic shifted-window cross-region "
        "communication. These architectural innovations significantly improve efficiency and have "
        "demonstrated strong performance on dense prediction tasks, yet they impose a fixed spatial "
        "attention topology that is data-agnostic. DAT [7] addressed this limitation through "
        "deformable attention, dynamically predicting sampling offsets from input content to attend "
        "to task-relevant spatial locations. Despite these advances in architecture design, a separate "
        "but equally important class of problems concerns the quality and interpretability of dense "
        "feature representations that ViT produces.")

    add_heading(doc, "Artifacts and Pathologies in ViT Dense Features", 3)
    add_text(doc,
        "While ViT excels at global image-level tasks, a growing body of evidence reveals systematic "
        "pathologies in its dense (patch-level) feature maps. Darcet et al. [8] first documented the "
        "high-norm token phenomenon: in self-supervised ViTs such as DINOv2, certain background tokens "
        "inflate to extreme activation norms, creating spatial artifacts that degrade dense prediction "
        "quality. They proposed adding register tokens—unused positions that absorb background "
        "noise—as a pragmatic fix. While effective at eliminating anomalous high-norm patches, "
        "register tokens treat the symptom without addressing the underlying cause, and do not improve "
        "or may even reduce point-in-box localization accuracy.")
    add_text(doc,
        "Concurrently, label-supervised ViTs exhibit attention deficit: the [CLS] attention map often "
        "disperses across irrelevant background regions rather than focusing on class-relevant objects [9]. "
        "In text-supervised models like CLIP, patch features may not align well with grounded textual "
        "semantics, leading to poor performance on open-vocabulary segmentation and detection tasks [10], "
        "[11]. These seemingly distinct pathologies share a common root: the coarse-grained supervision "
        "signals (image-level labels or image-caption pairs) do not constrain the spatial specificity of "
        "the aggregated representation.")

    add_heading(doc, "Diagnostic Frameworks for Patch-Level Quality", 3)
    add_text(doc,
        "A critical contribution toward understanding these pathologies was made by Shi et al. [12], "
        "who proposed two complementary diagnostic metrics. The Patch Score, defined as the cosine "
        "similarity between each patch feature and the global [CLS] representation, quantifies how "
        "strongly each spatial location contributes to the final image-level representation. The "
        "Point-in-Box (PiB) metric evaluates whether the highest-scoring patch falls within the "
        "annotated object bounding box. Using these metrics, Shi et al. demonstrated that vanilla "
        "ViT consistently assigns higher Patch Scores to background patches than foreground ones—"
        "achieving PiB values of only 19–48% compared to 60–80% for convolutional networks. "
        "Crucially, they showed that this low PiB emerges early in training and persists throughout "
        "optimization, confirming it as a structural pathology rather than a training artifact.")

    # ── 2.2 ──
    add_heading(doc, "Selective CLS Aggregation: From Frequency Stability to Adaptive Selection", 2)

    add_heading(doc, "LaSt-ViT: Frequency Stability as a Foreground Proxy", 3)
    add_text(doc,
        "Shi et al. [12] diagnosed the root cause of ViT dense feature artifacts as Lazy Aggregation: "
        "the [CLS] token, under image-level supervision, tends to aggregate semantic information from "
        "whichever patches are most easily accessible—typically the majority background patches—"
        "rather than from the semantically informative foreground patches. This occurs because background "
        "patches can absorb foreground information through global attention, and aggregating from many "
        "background patches is a lower-entropy solution that minimizes training loss.")
    add_text(doc,
        "To address Lazy Aggregation, Shi et al. proposed LaSt-ViT (Lazy Strike), which replaces the "
        "default all-patch aggregation with a frequency-based selection mechanism. For each patch "
        "feature, LaSt-ViT applies a 1D FFT along the channel dimension, passes the result through a "
        "Gaussian low-pass filter, and inverse-transforms back. The rationale is that foreground object "
        "patches exhibit more stable (low-frequency) channel-wise activations, while background patches "
        "tend toward higher-frequency, less coherent activations. A channel-wise Top-K pooling then "
        "selects the K most stable patches for each feature channel, and their features are averaged "
        "to form the aggregated representation.")
    add_text(doc,
        "While LaSt-ViT represents an important advance, it has several limitations that motivate our "
        "work. First, the fixed Top-K ratio creates a rigid foreground budget that cannot adapt to "
        "images with varying object sizes or complexity. Second, frequency stability alone may conflate "
        "smooth foreground objects (e.g., a white swan) with smooth background regions (e.g., sky or "
        "walls), leading to the stable-background failure mode that we systematically identify and "
        "address. Third, LaSt-ViT does not explicitly model the depth dimension of ViT's multi-layer "
        "representations, missing the opportunity to selectively aggregate information across layers.")

    add_heading(doc, "Adaptive and Multi-Cue Aggregation", 3)
    add_text(doc,
        "Several concurrent works have explored alternatives to fixed aggregation strategies. TokenCut "
        "[13] leverages self-similarity graphs and normalized cuts to segment foreground tokens without "
        "any training, achieving strong object discovery results. However, graph-based methods incur "
        "non-trivial computational overhead on large token sets. MaskCLIP [11] modifies CLIP's attention "
        "to produce more spatially grounded patch features through text-guided masking, but requires "
        "text supervision at inference time. FreeDA [14] proposes training-free debiased adaptation "
        "for CLIP, using covariance estimation to suppress background biases, yet operates at the "
        "output feature level rather than modifying the aggregation mechanism.")
    add_text(doc,
        "From a complementary direction, several works have explored adaptive token selection for "
        "efficiency. DynamicViT [15] and ATS [16] learn to prune unimportant tokens progressively "
        "through the network, improving inference speed but not specifically targeting the CLS "
        "aggregation bias. ToMe [17] merges similar tokens to reduce computation while preserving "
        "visual fidelity. These works demonstrate that intelligent spatial selection improves ViT "
        "representations, but they focus on efficiency rather than on correcting the fundamental "
        "CLS aggregation pathology.")
    add_text(doc,
        "Our work builds upon LaSt-ViT's insight of frequency-guided selection while extending it in "
        "three critical dimensions: (i) we enrich the foreground signal with multiple complementary "
        "cues beyond frequency stability alone; (ii) we replace the fixed Top-K with an image-adaptive "
        "soft mask predicted from the global image feature; and (iii) we decouple the foreground and "
        "background streams rather than simply suppressing background, enabling each stream to "
        "independently develop task-relevant semantics.")

    # ── 2.3 ──
    add_heading(doc, "Depth-Wise Information Aggregation in Deep Transformers", 2)

    add_heading(doc, "Residual Connections as Implicit Aggregators", 3)
    add_text(doc,
        "The residual connection, introduced by He et al. [18] for deep convolutional networks and "
        "adopted universally in Transformer architectures, serves a dual role that is often "
        "underappreciated. Beyond providing an identity path for stable gradient flow, residual "
        "connections implicitly define how information accumulates across depth. In the PreNorm "
        "formulation standard in modern Transformers, the hidden state at layer l is computed as "
        "h_l = h_{l-1} + f_l(h_{l-1}). Unrolling this recursion reveals that the input to any "
        "layer is the sum of the initial embedding and all previous layer outputs, each weighted "
        "uniformly by 1.")
    add_text(doc,
        "This uniform accumulation creates a fundamental tension in deep Transformers. As network "
        "depth increases, the hidden state magnitude grows approximately as O(L), causing the "
        "relative contribution of any single layer to be progressively diluted—a phenomenon "
        "termed PreNorm dilution [19]. Early-layer features encoding low-level patterns become "
        "entangled in a monolithic representation, and later layers cannot selectively retrieve "
        "specific earlier information. This problem is particularly acute in Vision Transformers, "
        "where different layers capture fundamentally different levels of abstraction: early layers "
        "encode textures and edges, middle layers capture parts and patterns, and later layers "
        "encode object-level semantics [20], [21].")

    add_heading(doc, "Gated and Selective Residual Connections", 3)
    add_text(doc,
        "Recognizing the limitations of fixed residual addition, several works have introduced "
        "learnable gating mechanisms. Highway Networks [22] pioneered this direction by adding "
        "data-dependent gates that control information flow through layers. DenseNet [23] adopted "
        "an alternative strategy of concatenating rather than adding features from all preceding "
        "layers, giving later layers explicit access to all earlier representations at the cost "
        "of linearly growing memory.")
    add_text(doc,
        "In the Transformer context, GPT-2 [24] introduced learnable scaling factors on residual "
        "paths, initialized to 1/sqrt(N) where N is the number of residual connections, to "
        "stabilize training of deep models. More recently, DenseFormer [25] extended this idea "
        "by learning per-layer static weights for all historical representations, providing "
        "content-independent but layer-specific aggregation. While an improvement over uniform "
        "addition, static weights cannot adapt to input-specific information needs. SubLN [26] "
        "revisited PreNorm residual scaling and demonstrated that careful normalization within "
        "residual branches can partially mitigate the dilution effect.")

    add_heading(doc, "Attention Residuals: Content-Dependent Depth Aggregation", 3)
    add_text(doc,
        "The most directly relevant prior work on depth-wise aggregation is Attention Residuals "
        "(AttnRes) [27], which fundamentally reconceptualizes the residual connection as a "
        "depth-dimension attention mechanism. In AttnRes, each layer l receives its input not "
        "through uniform summation but through a learned softmax attention over all historical "
        "layer outputs. The attention weights are computed using a learned pseudo-query vector "
        "as the query and RMSNorm-transformed historical outputs as keys, enabling content-dependent "
        "depth selection.")
    add_text(doc,
        "AttnRes demonstrated that this selective depth aggregation yields a 1.25× compute "
        "advantage on language modeling benchmarks. The authors further proposed Block Attention "
        "Residuals, which partitions layers into blocks and applies inter-block attention, "
        "reducing storage from O(Ld) to O(Nd) for N blocks while preserving most of the benefits. "
        "Empirical analysis on a 48B-parameter model showed that Block AttnRes eliminates the "
        "monotonic growth of hidden-state magnitude and produces more uniform gradient distributions "
        "across depth.")
    add_text(doc,
        "While AttnRes was developed in the context of language models, its core insight—that "
        "depth-wise aggregation should be content-adaptive rather than fixed—is directly "
        "applicable to Vision Transformers. However, AttnRes addresses all tokens uniformly "
        "through a single aggregation mechanism, without distinguishing between foreground and "
        "background information streams. Our method extends this paradigm by coupling depth-wise "
        "attention with foreground-background decoupling, enabling independent depth aggregation "
        "for each semantic stream.")

    # ── 2.4 ──
    add_heading(doc, "Foreground-Background Separation in Visual Representation Learning", 2)

    add_heading(doc, "Class Activation Mapping and CAM-Based Methods", 3)
    add_text(doc,
        "Separating foreground from background is a longstanding challenge in computer vision. "
        "Class Activation Mapping (CAM) [28] and its extensions—GradCAM [29], GradCAM++ [30], "
        "and ScoreCAM [31]—generate coarse localization heatmaps from classification networks "
        "by back-propagating or weighting activation maps. These methods have proven invaluable "
        "for weakly supervised localization and segmentation, but they inherit the same background "
        "bias as the underlying classifier. DINO-seg [32] extracts self-attention maps from "
        "self-supervised ViTs to produce unsupervised object segmentations, but its attention "
        "maps still suffer from background activations and require post-processing that introduces "
        "additional hyperparameters.")

    add_heading(doc, "Multi-Cue Visual Saliency and Objectness", 3)
    add_text(doc,
        "Visual saliency detection has a rich history of leveraging multiple complementary cues. "
        "Classic methods combine color contrast, texture uniqueness, and spatial priors to predict "
        "fixation points [33], [34]. The principle that no single cue is sufficient for robust "
        "saliency detection—and that cues should be adaptively weighted based on context—"
        "directly motivates our multi-cue foregroundness approach. In the ViT feature space "
        "specifically, TokenCut [13] uses self-similarity as a graph-theoretic proxy for "
        "objectness. LOST [35] combines self-similarity with a seeding strategy that progressively "
        "refines the foreground region. FreeSeg [36] segments images without training by combining "
        "feature similarity with spatial compactness. These methods demonstrate that combining "
        "multiple signals yields more robust foreground estimates than any single measure, but they "
        "operate at the feature level without modifying the underlying aggregation mechanism.")

    # ── 2.5 ──
    add_heading(doc, "Connections to Post-Hoc Aggregation and Frozen Backbone Methods", 2)
    add_text(doc,
        "An important design principle that our work shares with several prior methods is operating "
        "in a post-hoc manner on frozen pretrained backbones. This paradigm, exemplified by methods "
        "like MaskCLIP [11], FreeDA [14], and Saliency-guided Open-Vocabulary Segmentation [37], "
        "offers several practical advantages: it avoids costly retraining of large foundation models, "
        "preserves the backbone's generalization properties, and allows the aggregation module to be "
        "lightweight and fast to train. Our approach follows this paradigm, training only approximately "
        "0.59M parameters (2.7% of the DINO ViT-S/16 backbone) while achieving substantial "
        "improvements in dense feature quality.")
    add_text(doc,
        "More broadly, our work fits into the emerging trend of studying and correcting information "
        "flow within pretrained Transformers, rather than simply scaling up model or data size. Just "
        "as Attention Residuals revealed that the default residual connection is suboptimal for "
        "information aggregation across depth, we show that the default CLS pooling is suboptimal "
        "for spatial aggregation. Both findings point to the same overarching lesson: the default "
        "aggregation operations in Transformers—whether across layers or across spatial "
        "positions—are overly rigid and fail to exploit the structured information available "
        "in the intermediate representations.")

    # ── 2.6 ──
    add_heading(doc, "Summary: Positioning Our Work", 2)
    add_text(doc,
        "TABLE I positions our work relative to the most closely related methods along four key "
        "dimensions: spatial selectivity, multi-cue foregroundness, adaptive budget, and "
        "depth-wise attention. LaSt-ViT introduced spatial selectivity through frequency stability "
        "but uses fixed Top-K with a single cue. Attention Residuals introduced depth-wise "
        "attention but operates uniformly across spatial positions. Our DAFB-CLS framework "
        "is the first to unify all four dimensions into a single coherent framework for CLS "
        "aggregation in Vision Transformers.")
    add_text(doc, "TABLE I", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=0)
    add_text(doc, "COMPARISON WITH RELATED METHODS", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=0)
    add_table(doc,
        ["Method", "Spatial Select.", "Multi-Cue FG", "Adapt. Budget", "Depth Attn"],
        [
            ["CAM / GradCAM",   "No",  "No",          "No",           "No"],
            ["DINO-seg",        "Partial", "No",      "No",           "No"],
            ["TokenCut",        "Yes", "Partial",     "No",           "No"],
            ["LaSt-ViT",        "Yes", "No (freq)",   "No (fixed K)", "No"],
            ["Attn Residuals",  "No",  "No",          "N/A",          "Yes"],
            ["DenseFormer",     "No",  "No",          "N/A",          "Partial"],
            ["DAFB-CLS (Ours)", "Yes", "Yes (4 cues)","Yes (learned)","Yes (F/B)"],
        ])


def write_methods(doc):
    add_heading(doc, "PROPOSED METHOD", 1)

    add_heading(doc, "Problem Formulation and Design Rationale", 2)
    add_text(doc,
        "We consider a frozen Vision Transformer backbone (e.g., DINO ViT-S/16 or OpenCLIP ViT-B/16) "
        "that processes an image x of size H×W×3 into a sequence of N = (H/P)×(W/P) "
        "patch tokens plus a [CLS] token, where P is the patch size. The backbone consists of L "
        "transformer blocks; we extract intermediate features from a subset of layers "
        "λ = {l_1, l_2, ..., l_K} (typically K=4 layers equally spaced, e.g., {3, 6, 9, 12} "
        "for a 12-layer ViT).")
    add_text(doc,
        "Our goal is to train a lightweight post-hoc aggregation module that transforms these "
        "multi-layer features into an improved image representation C, without modifying or "
        "retraining the frozen backbone. This design is motivated by two complementary insights: "
        "(i) the default [CLS] pooling suffers from Lazy Aggregation bias toward background "
        "tokens [12], and (ii) the standard residual connections accumulate layer outputs with "
        "fixed uniform weights, lacking content-adaptive depth selection [27].")
    add_text(doc,
        "Our key design principle is to address both problems jointly through foreground-background "
        "decoupling: rather than suppressing background information, we maintain two separate "
        "semantic streams—a foreground stream C_F that aggregates foreground-identifying "
        "features, and a background stream C_B that aggregates contextual scene features—and "
        "learn to fuse them adaptively based on the task and input.")

    add_heading(doc, "Step 1: Multi-Layer Feature Extraction", 2)
    add_text(doc,
        "Given input image x, we first pass it through the frozen ViT backbone and extract "
        "intermediate patch features via forward hooks registered at the specified layers λ. "
        "Let P^l ∈ R^{B×N×D} denote the patch token features (excluding [CLS]) "
        "at layer l, where B is the batch size, N is the number of patches, and D is the feature "
        "dimension. We stack features across layers to form a 4D tensor "
        "Π ∈ R^{B×K×N×D}. The [CLS] token feature from the final "
        "extracted layer provides the global image representation c_cls ∈ R^{B×D}. "
        "All features are detached from the backbone computation graph to ensure that backbone "
        "parameters remain frozen during training.")

    add_heading(doc, "Step 2: Multi-Cue Foregroundness Scoring", 2)
    add_text(doc,
        "The first major component computes a per-patch foregroundness score that indicates how "
        "likely each patch is to belong to a foreground object. Motivated by the observation that "
        "no single signal reliably distinguishes foreground from background across diverse image "
        "types, we design four complementary cues that capture different aspects of foregroundness.")

    add_heading(doc, "Frequency Stability Cue", 3)
    add_text(doc,
        "Inspired by LaSt-ViT, we measure the stability of each patch's feature vector under "
        "spectral low-pass filtering. For each patch feature p_i ∈ R^D, we compute its 1D FFT "
        "along the channel dimension, apply a Gaussian low-pass filter with cutoff frequency "
        "σ_f = 0.25, and inverse-transform to obtain the filtered version p̃_i. The "
        "stability score for patch i is S_i = (1/D)∑_d p̃_{i,d}/(|p̃_{i,d} - "
        "p_{i,d}| + ε). Higher values indicate that the patch feature changes minimally under "
        "spectral filtering, suggesting stable, low-frequency semantic content typical of "
        "foreground objects.")

    add_heading(doc, "Depth Consistency Cue", 3)
    add_text(doc,
        "While the frequency cue captures within-layer stability, the depth consistency cue "
        "captures stability across layers. Patches corresponding to semantically meaningful objects "
        "tend to maintain more consistent representations across network depth, while background "
        "patches exhibit greater inter-layer variation. We measure depth consistency as the average "
        "cosine similarity between each layer's patch feature and the cross-layer mean: "
        "D_i = (1/K)∑_k cos(P_i^{l_k}, p̄_i), where p̄_i = (1/K)∑_k P_i^{l_k} "
        "is the mean feature across layers.")

    add_heading(doc, "Semantic Alignment Cue", 3)
    add_text(doc,
        "The semantic alignment cue measures how strongly each patch's feature aligns with the "
        "global image semantics: A_i = (1/K)∑_k cos(P_i^{l_k}, c_cls). For text-supervised "
        "backbones (e.g., OpenCLIP), we additionally support text-guided semantic alignment, "
        "where the similarity is computed between each patch and the set of text embeddings, "
        "using the maximum similarity across all text embeddings.")

    add_heading(doc, "Spatial Compactness Cue", 3)
    add_text(doc,
        "Foreground objects in natural images typically occupy spatially contiguous regions. The "
        "spatial compactness cue enforces this prior by applying local neighborhood smoothing to "
        "the semantic alignment scores via average pooling with kernel size k_s = 3, encouraging "
        "patches within the same spatial neighborhood to have similar foregroundness. This cue does "
        "not introduce additional learnable parameters but provides a useful structural inductive "
        "bias.")

    add_heading(doc, "Learned Cue Fusion", 3)
    add_text(doc,
        "The four cues are fused through a combination of learnable weighting and an MLP "
        "refinement module. Let s_i = [S_i, D_i, A_i, C_i]^T denote the cue vector for patch i. "
        "We first compute a weighted combination with learned softmax-normalized weights: "
        "F_i^linear = ∑_k w_k · s_{i,k}, where w = softmax(θ_w) are learnable "
        "parameters initialized uniformly. In parallel, a lightweight MLP with two hidden layers "
        "(hidden dimension 256, GELU activations) refines the cue vector to capture nonlinear "
        "interactions: F_i^mlp = MLP(s_i). The final foregroundness score combines both components "
        "and applies spatial smoothing via a learnable 3×3 convolution initialized to "
        "uniform 1/9: F_i = F_i^linear + F_i^mlp + 0.5 · σ(W_s * F_spatial). "
        "The total ForegroundnessHead introduces fewer than 5K parameters.")

    add_heading(doc, "Step 3: Adaptive Foreground Budget", 2)
    add_text(doc,
        "A critical design decision is how to convert the continuous foregroundness scores into "
        "foreground masks. Existing methods typically use fixed-ratio Top-K selection, where the "
        "top K% of patches are selected regardless of image content. This rigid budget is "
        "problematic: images with a single small object may have fewer than 10% foreground patches, "
        "while images with multiple large objects may have over 50%.")
    add_text(doc,
        "We address this with an Adaptive Budget Module that predicts an image-specific threshold "
        "τ from the global feature c_cls, and applies a temperature-controlled sigmoid to "
        "produce a soft foreground mask: m_i^F = σ((F_i - τ)/T), where τ = "
        "MLP_τ(c_cls) is a two-layer MLP mapping the global feature to a scalar threshold, "
        "and T is a learnable temperature parameter initialized to 0.1 (in log-space for stable "
        "optimization) and clamped to [0.01, 1.0].")
    add_text(doc,
        "The adaptive threshold mechanism has several advantages over fixed Top-K: (i) it naturally "
        "adapts to images with different foreground proportions; (ii) the soft mask avoids the "
        "discontinuity of hard thresholding, which creates artifacts at mask boundaries and "
        "destabilizes training gradients; and (iii) the learnable temperature enables the model to "
        "control mask sharpness as a function of training progress.")

    add_heading(doc, "Step 4: Dual CLS Decoupling", 2)

    add_heading(doc, "Foreground CLS Aggregation", 3)
    add_text(doc,
        "Given the soft foreground mask m^F, we compute the foreground CLS representation at each "
        "extracted layer by weighted averaging of patch features: "
        "B_l^F = (∑_i m_i^F · P_i^l) / (∑_i m_i^F + ε). This produces a "
        "sequence of foreground CLS tokens B^F = [B_{l_1}^F, ..., B_{l_K}^F] ∈ R^{K×D} "
        "that represent the foreground content as seen from each extracted layer.")

    add_heading(doc, "Background CLS Aggregation", 3)
    add_text(doc,
        "Rather than simply using m^B = 1 - m^F as the background mask, we introduce a dedicated "
        "BackgroundnessHead that independently predicts a background mask from the final-layer "
        "patch features. This head consists of a two-layer MLP that outputs a background score, "
        "combined with an uncertainty predictor that estimates the reliability of each patch's "
        "backgroundness prediction: m_i^B = s_i^B · (1 - m_i^F) · u_i, where s_i^B = "
        "σ(MLP_B(P_i^{l_K})) is the background score and u_i = σ(MLP_U(P_i^{l_K})) is "
        "the uncertainty estimate. The multiplicative composition ensures that high-confidence "
        "foreground patches are excluded, uncertain boundary patches are down-weighted, and the "
        "background stream can develop its own semantic preferences.")

    add_heading(doc, "Why Dual Decoupling Rather Than Background Suppression", 3)
    add_text(doc,
        "Three reasons motivate maintaining a separate background stream rather than simply "
        "suppressing background information. First, background context provides complementary "
        "semantic information: scene type, spatial layout, and co-occurrence statistics that are "
        "valuable for classification and may aid localization. Second, completely suppressing "
        "background forces all contextual information through the foreground stream, which may "
        "dilute the foreground signal. Third, maintaining separate streams enables the "
        "task-adaptive fusion to learn how much background context each task needs, which varies "
        "substantially: object discovery benefits from minimal background, while classification "
        "may benefit from scene context.")

    add_heading(doc, "Step 5: Depth-Wise Attention", 2)
    add_text(doc,
        "After obtaining the dual CLS sequences B^F and B^B, we perform independent depth-wise "
        "attention over the layer dimension for each stream. For the foreground stream: "
        "β_l^F = softmax(w_F^T · RMSNorm(B_l^F)), C_F = ∑_l β_l^F · "
        "B_l^F, where w_F ∈ R^D is a learned pseudo-query vector initialized to zero, "
        "ensuring uniform initial attention weights. RMSNorm is applied to the keys (block "
        "features) to prevent layers with large activation magnitudes from dominating the "
        "attention weights—a key design choice validated by the AttnRes ablation study [27]. "
        "The background stream uses an analogous formulation with its own pseudo-query w_B.")
    add_text(doc,
        "The zero initialization of pseudo-queries is crucial: at the start of training, all "
        "layers contribute equally (β_l = 1/K), so the initial behavior matches standard "
        "average pooling across layers. As training progresses, the model learns to emphasize "
        "layers that are most informative for each stream—typically later layers for semantic "
        "understanding and earlier layers for spatial precision.")

    add_heading(doc, "Step 6: Task-Adaptive Fusion", 2)
    add_text(doc,
        "The final image representation is produced by adaptively fusing the foreground and "
        "background streams: g = σ(MLP_g([C_F; C_B; c_cls])), C = g · C_F + (1 - g) "
        "· C_B, where g ∈ [0, 1] is a learnable gate predicted by a two-layer MLP "
        "mapping from 3D to 1. The gate takes all three representations as input, allowing it to "
        "make an informed decision about how much foreground vs. background information the task "
        "requires. For example, object discovery tasks may learn g → 1, while classification "
        "tasks may learn a more balanced gate value.")
    add_text(doc,
        "The fused representation C is then passed to the appropriate task head: for "
        "classification, a two-layer MLP maps C to class logits; for segmentation, patch features "
        "from the final layer are projected to the text embedding space (for OpenCLIP) or the CLS "
        "projection space, producing per-patch similarity scores reshaped to spatial maps; for "
        "object discovery, a scoring MLP produces per-patch scores combined with the foreground "
        "mask and cosine similarity to C.")

    add_heading(doc, "Training Objective", 2)
    add_text(doc,
        "The total training loss combines a task-specific primary loss with three auxiliary losses: "
        "L = L_task + λ_fg L_mask + λ_dec L_decouple + λ_bud L_budget.")

    add_heading(doc, "Task Loss", 3)
    add_text(doc,
        "For object discovery, we supervise the foreground mask using pseudo-ground-truth masks "
        "derived from the initial foregroundness scores. Patches with scores exceeding one standard "
        "deviation above the mean are labeled as foreground, generating binary pseudo-masks. The "
        "task loss combines binary cross-entropy and Dice loss for robust segmentation supervision: "
        "L_mask = BCE(m^F, m̂) + (1 - Dice(m^F, m̂)).")

    add_heading(doc, "Decoupling Loss", 3)
    add_text(doc,
        "To encourage the foreground and background streams to capture complementary rather than "
        "redundant information, we minimize the squared cosine similarity between C_F and C_B: "
        "L_decouple = (cos(C_F, C_B))^2. The squared formulation provides stronger gradients "
        "when the streams are nearly orthogonal, preventing the loss from becoming ineffective "
        "late in training. The loss weight is λ_dec = 0.05.")

    add_heading(doc, "Budget Regularization", 3)
    add_text(doc,
        "Without regularization, the adaptive budget may converge to degenerate solutions. We "
        "prevent these failure modes through a range penalty on the average foreground ratio "
        "r = (1/N)∑_i m_i^F: L_budget = max(0, r_min - r)^2 + max(0, r - r_max)^2, "
        "where r_min = 0.1 and r_max = 0.7 are the lower and upper bounds. The loss weight is "
        "λ_bud = 0.01.")

    add_heading(doc, "Complexity Analysis", 2)
    add_text(doc,
        "The DAFB-CLS aggregation module is designed to be lightweight. For a standard "
        "ViT-S/16 with D = 384, N = 196, and K = 4 extracted layers, the total trainable "
        "parameters are 0.59M (DINO ViT-S/16) or 1.74M (OpenCLIP ViT-B/16), representing "
        "2.7% and 2.0% of the respective frozen backbones. Inference latency overhead is "
        "1.82ms (DINO) and 3.45ms (OpenCLIP), both under 31% relative overhead. Peak GPU "
        "memory scales linearly with the number of extracted layers, resulting in approximately "
        "2.8× (DINO) and 3.0× (OpenCLIP) overhead relative to the baseline ViT "
        "forward pass.")


def main():
    doc = Document(TEMPLATE)
    clear_doc(doc)

    # Title
    p = doc.add_paragraph(style="Title")
    run = p.add_run("DAFB-CLS: Depth-Adaptive Foreground-Background CLS Decoupling for Vision Transformers")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Author placeholder
    p = doc.add_paragraph(style="Authors")
    run = p.add_run("Author Name, Member, IEEE")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Abstract
    p = doc.add_paragraph(style="Abstract")
    run = p.add_run(
        "Abstract—Vision Transformers (ViTs) have achieved remarkable success across visual "
        "recognition tasks, yet the default [CLS] token aggregation suffers from Lazy Aggregation: "
        "the tendency to aggregate semantics from background patches as shortcuts under coarse-grained "
        "image-level supervision. Meanwhile, standard residual connections accumulate layer outputs "
        "with fixed uniform weights, lacking content-adaptive depth selection. We propose DAFB-CLS, "
        "a lightweight post-hoc aggregation framework that addresses both problems jointly through "
        "foreground-background decoupling. DAFB-CLS introduces multi-cue foregroundness scoring "
        "(combining frequency stability, depth consistency, semantic alignment, and spatial "
        "compactness), an adaptive soft foreground mask predicted from global image features, dual "
        "CLS streams with independent depth-wise attention inspired by Attention Residuals, and "
        "task-adaptive fusion. Training only 0.59M parameters (2.7% of the frozen DINO ViT-S/16 "
        "backbone), DAFB-CLS improves CorLoc by +50pp over the raw ViT and +13.7pp over LaSt-ViT "
        "on VOC, with Mask IoU gains of +18pp. Consistent improvements are observed on COCO "
        "(+11.4pp CorLoc, +15.8pp Mask IoU) and across OpenCLIP ViT-B/16 segmentation tasks.")

    # Index Terms
    p = doc.add_paragraph(style="IndexTerms")
    run = p.add_run(
        "Index Terms—Vision Transformer, CLS token aggregation, foreground-background "
        "decoupling, depth-wise attention, object discovery, dense prediction.")

    write_related_work(doc)
    doc.add_page_break()
    write_methods(doc)

    doc.save(OUTPUT)
    print(f"Saved to: {OUTPUT}")

if __name__ == "__main__":
    main()
