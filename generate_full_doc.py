"""
生成DAFB-CLS完整项目详细文档（一份文档完全了解项目）
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Microsoft YaHei'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    hs.font.color.rgb = RGBColor(0, 51, 102)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
    return table

def bp(doc, text):
    doc.add_paragraph(text, style='List Bullet')

def add_formula(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(80, 80, 80)

def bold_para(doc, bold_text, normal_text=''):
    p = doc.add_paragraph()
    r = p.add_run(bold_text)
    r.bold = True
    if normal_text:
        p.add_run(normal_text)

# ==================== 封面 ====================
title = doc.add_heading('DAFB-CLS: Depth-Adaptive Foreground-Background CLS Decoupling for Vision Transformers', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('完整项目文档')
r.font.size = Pt(16)
r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('研究方向：计算机视觉 > 视觉Transformer > 可解释性与空间定位')

doc.add_page_break()

# ==================== 目录概览 ====================
doc.add_heading('目录概览', level=1)
for item in [
    '一、项目概述',
    '二、研究背景与动机',
    '三、相关工作',
    '四、方法（DAFB-CLS）',
    '五、实验设置',
    '六、实验结果',
    '七、消融实验',
    '八、压力测试与失败模式分析',
    '九、改进尝试与未来方向',
    '十、代码架构说明',
    '十一、完整参考文献',
]:
    doc.add_paragraph(item)

doc.add_page_break()

# ==================== 一、项目概述 ====================
doc.add_heading('一、项目概述', level=1)

doc.add_paragraph(
    'DAFB-CLS（Depth-Adaptive Foreground-Background CLS Decoupling for Vision Transformers）'
    '是一个针对Vision Transformer（ViT）中CLS token质量问题的后处理聚合框架。'
    '项目开发时间：2026年5月11日-18日。'
)

doc.add_paragraph('项目的核心贡献：')
bp(doc, '提出DAFB-CLS框架，将CLS token解耦为独立的前景和背景语义流')
bp(doc, '设计多线索前景区分机制，融合四种互补线索')
bp(doc, '设计自适应软掩码，替代固定Top-K选择')
bp(doc, '设计前景/背景独立的深度注意力')
bp(doc, '在VOC和COCO两个数据集上验证了方法的有效性和泛化性')

doc.add_paragraph('项目的关键数据：')
add_table(doc,
    ['数据集', '骨干', 'CorLoc', 'Mask IoU', 'vs LaSt-ViT CorLoc', 'vs LaSt-ViT Mask IoU'],
    [
        ['VOC 2012', 'DINO ViT-S/16', '79.43%', '31.27%', '+13.73pp', '+17.96pp'],
        ['COCO 2017', 'DINO ViT-S/16', '72.96%', '28.73%', '+11.43pp', '+15.78pp'],
    ]
)

doc.add_page_break()

# ==================== 二、研究背景与动机 ====================
doc.add_heading('二、研究背景与动机', level=1)

doc.add_heading('2.1 Vision Transformer中的CLS Token', level=2)
doc.add_paragraph(
    'Vision Transformer（ViT）将图像分割为固定大小的patch序列（默认14x14=196个patch），'
    '每个patch被编码为D维向量（DINO ViT-S/16为384维，OpenCLIP ViT-B/16为768维）。'
    '一个特殊的CLS token被添加到序列开头，通过L层Transformer的自注意力机制与所有patch交互，'
    '最终汇聚为一个D维的全局表示向量，用于下游任务（分类、分割、物体发现等）。'
)

doc.add_heading('2.2 问题一：惰性聚合（Lazy Aggregation）', level=2)
doc.add_paragraph(
    '该问题由Darcet等在ICLR 2024论文"Vision Transformers Need More Than Registers"中首次系统性揭示。'
    '研究发现，CLS token在聚合信息时存在两个问题：'
)
bp(doc, 'Type-1 artifact：单一patch接收过多注意力权重，形成注意力"黑洞"')
bp(doc, 'Type-2 artifact：注意力在所有token上均匀分布，导致语义信息"蒸发"')
doc.add_paragraph(
    '这些artifact token本质上充当了"垃圾桶"角色，CLS token将不需要的注意力权重倾倒在这些位置，'
    '从而绕过了对前景区域的精确建模。改进版方案提出用前景/背景register替代artifact token，'
    '但需要重新训练模型，无法即插即用。'
)

doc.add_heading('2.3 问题二：惰性累加（Lazy Accumulation）', level=2)
doc.add_paragraph(
    '该问题由"Attention Residuals"论文指出。标准残差连接形式为 x + F(x)，'
    '即每一层的输出与输入进行等权重累加。当网络有L层时，最终输出是所有历史层输出的等权重叠加。'
    '这意味着模型无法自适应地决定哪些层的输出对当前任务最重要。'
)

doc.add_heading('2.4 现有方法的不足', level=2)
add_table(doc,
    ['方法', '优势', '不足'],
    [
        ['Register (ICLR 2024)', '简洁有效，不改变推理', '需重新训练，未区分前景/背景'],
        ['Attention Residuals', '深度自适应聚合', '单流操作，未区分前景/背景'],
        ['LaSt-ViT', '频率稳定性信号', '单一信号，固定Top-K选择'],
        ['CAM/GradCAM', '梯度可视化', '缺乏前/背景显式建模'],
        ['DINO-seg', '自注意力图分割', '单一信号，阈值敏感'],
    ]
)

doc.add_page_break()

# ==================== 三、相关工作 ====================
doc.add_heading('三、相关工作', level=1)

doc.add_heading('3.1 Vision Transformers', level=2)
doc.add_paragraph(
    'ViT [1]首次将Transformer引入视觉领域。DeiT [2]通过知识蒸馏改善训练效率。'
    'DINO [3]和DINOv2 [4]通过自蒸馏学习高质量视觉表示。'
    'CLIP [5]和OpenCLIP [6]通过对比学习实现视觉-语言对齐。'
    'MoCo v3 [13]将对比学习引入ViT训练，iBOT [14]结合实例判别与掩码图像建模，'
    'BEiT [11]和MAE [12]分别通过离散token预测和掩码自编码进行预训练。'
)

doc.add_heading('3.2 注意力机制与可解释性', level=2)
doc.add_paragraph(
    'CAM [7]利用全局平均池化权重生成激活图。GradCAM [8]通过梯度加权激活图。'
    'ScoreCAM [29]提出无梯度版本。Attention Rollout [9]递归聚合所有层注意力权重。'
    'DINO-seg [10]利用DINO自注意力图的阈值化获得分割掩码。'
    'TokenCut [20]将自注意力图建模为图结构，通过归一化割实现无监督物体分割。'
)

doc.add_heading('3.3 Register Token', level=2)
doc.add_paragraph(
    'Darcet等 [15]首次揭示DINOv2的artifact现象，提出在训练时添加register token消除artifact。'
    '改进版 [16]进一步将artifact register替换为前景/背景register。'
    '但该方法需重新训练模型，register token在所有层共享，缺乏深度选择性。'
)

doc.add_heading('3.4 残差连接与深度聚合', level=2)
doc.add_paragraph(
    'Transformer [32]的残差连接对所有层等权叠加。DeepViT [18]发现深层自注意力趋于相似（注意力坍缩）。'
    'Attention Residuals [17]提出用注意力机制替代固定累加，但仍在单一混合CLS上操作。'
)

doc.add_heading('3.5 物体发现与开放词汇分割', level=2)
doc.add_paragraph(
    'LOST [19]利用DINO自注意力图定位物体。FreeSOLO [21]通过自监督学习预测实例掩码。'
    'LaSt-ViT [22]利用FFT频率稳定性区分前景和背景patch。'
    'MaskCLIP [24]利用CLIP的patch-text相似度生成掩码提议。'
    'OVSeg [25]通过微调CLIP改善开放词汇分割。'
    'GroupViT [27]通过分组机制学习层次化视觉token。'
)

doc.add_page_break()

# ==================== 四、方法 ====================
doc.add_heading('四、方法（DAFB-CLS）', level=1)

doc.add_heading('4.1 总体框架', level=2)
doc.add_paragraph(
    'DAFB-CLS在冻结的预训练ViT骨干之上训练轻量聚合器，通过五个阶段将CLS token解耦为'
    '前景与背景语义流。给定输入图像（224x224），骨干网络通过前向钩子提取第3、6、9、12层的'
    'patch特征（形状[B, 4, 196, D]）与末层CLS token（形状[B, D]）。'
)

add_formula(doc, '阶段零：{x^l}_{l=1}^L, c = Backbone(I)')
add_formula(doc, '阶段一：F = CueFusion(x^L) → 前景性分数 [B, 196]')
add_formula(doc, '阶段二：m^F, m^B = MaskPrediction(F, c) → 前景/背景掩码 [B, 196]')
add_formula(doc, '阶段三：B_l^F = WeightedAvg(x^l, m^F), B_l^B = WeightedAvg(x^l, m^B)')
add_formula(doc, '阶段四：C_F = DepthAttn({B_l^F}), C_B = DepthAttn({B_l^B})')
add_formula(doc, '阶段五：C = g * C_F + (1-g) * C_B')

doc.add_heading('4.2 阶段一：多线索前景区分度计算', level=2)

doc.add_paragraph('四种互补线索：')

bold_para(doc, '线索一：频率稳定性（FrequencyStabilityCue）')
doc.add_paragraph(
    '对每层patch特征做FFT变换，施加截止频率sigma=0.25的高斯低通滤波器，'
    '计算滤波前后特征的残差比值。平滑背景（天空、墙面）因低频主导呈现高分，'
    '纹理前景（动物、建筑）因高频丰富呈现低分。'
)
add_formula(doc, 'x_fft = FFT(x), H(f) = exp(-0.5*(f/0.25)^2)')
add_formula(doc, 'x_filtered = IFFT(x_fft * H(f)).real')
add_formula(doc, 's_freq = mean_D(x_filtered / (|x_filtered - x| + 1e-6))')

bold_para(doc, '线索二：深度一致性（DepthConsistencyCue）')
doc.add_paragraph(
    '计算每个patch在4层的特征与跨层均值特征的余弦相似度。'
    '前景目标在不同层保持一致的激活，得分高；背景在不同层变化大，得分低。'
)
add_formula(doc, 'x_mean = mean_L(x_i^l)')
add_formula(doc, 's_depth = mean_L(cosine(x_i^l, x_mean))')

bold_para(doc, '线索三：语义对齐（SemanticAlignmentCue）')
doc.add_paragraph(
    '计算patch特征与CLS token的L2归一化余弦相似度，跨层平均。'
    '与全局语义更相关的patch更可能是前景。'
)
add_formula(doc, 's_sem = mean_L(cosine(x_i^l, c))')

bold_para(doc, '线索四：空间紧凑性（SpatialCompactnessCue）')
doc.add_paragraph(
    '对语义对齐分数的14x14二维重排施加3x3平均池化（步长1，填充1），'
    '抑制孤立高分patch，鼓励空间连续的前景区域。'
)
add_formula(doc, 's_spatial = AvgPool2d(reshape(s_sem, 14, 14), kernel=3)')

bold_para(doc, '线索融合方式')
doc.add_paragraph(
    '双重机制：（1）可学习softmax归一化权重（初始化各0.25）线性加权；'
    '（2）四条线索拼接为4维向量经3层MLP（4→256→256→1，GELU激活）非线性修正；'
    '（3）3x3卷积空间平滑残差（权重初始化1/9，固定系数0.5叠加，Sigmoid激活）。'
)

doc.add_heading('4.3 阶段二：自适应前景预算', level=2)
doc.add_paragraph(
    '阈值预测：两层MLP（D→128→1，GELU激活）从全局CLS token预测阈值tau。'
    '可学习温度T（初始0.1，约束[0.01, 1.0]）。'
    '软前景掩码：m^F = sigmoid((F - tau) / T)。'
)
doc.add_paragraph(
    '背景掩码：两个并行MLP（D→128→1）分别预测背景概率和不确定性估计。'
    'm^B = sigmoid(MLP_bg(x^L)) * (1 - m^F) * sigmoid(MLP_unc(x^L))。'
    '三项因子：背景概率、前景互补约束、不确定性门控。'
)

doc.add_heading('4.4 阶段三：双CLS解耦聚合', level=2)
doc.add_paragraph(
    '对每个提取层（第3、6、9、12层）独立执行前景-背景分离的归一化加权平均：'
)
add_formula(doc, 'B_l^F = sum_i(m_i^F * x_i^l) / sum_i(m_i^F)')
add_formula(doc, 'B_l^B = sum_i(m_i^B * x_i^l) / sum_i(m_i^B)')
doc.add_paragraph('加权分母clamp最小值1e-6确保数值稳定。在聚合源头阻断前景与背景语义交叉。')

doc.add_heading('4.5 阶段四：深度自适应注意力', level=2)
doc.add_paragraph(
    '为前景和背景流各自配备独立的深度注意力模块。'
    '每条流维护零初始化的query向量，对4层CLS序列经RMSNorm归一化（eps=1e-6，权重初始化全1）后作为key，'
    '通过内积计算深度注意力分数，softmax归一化后加权聚合。'
)
add_formula(doc, 'keys = RMSNorm(B^F)')
add_formula(doc, 'beta = softmax(K * q)')
add_formula(doc, 'C = sum(beta_l * B_l)')
doc.add_paragraph('两流参数完全独立，允许差异化的深度权重分布。')

doc.add_heading('4.6 阶段五：任务自适应融合', level=2)
doc.add_paragraph(
    '门控网络：MLP（3D→256→1，GELU激活，Sigmoid输出），输入为[C_F; C_B; c]拼接。'
    'C = g * C_F + (1-g) * C_B。逐图像预测标量门控值。'
)

doc.add_heading('4.7 损失函数', level=2)
add_formula(doc, 'L_total = L_task + lambda_fg * L_mask + lambda_dec * L_decouple + lambda_bud * L_budget')

add_table(doc,
    ['损失项', '公式', '作用', '权重(DINO/OpenCLIP)'],
    [
        ['L_task', '交叉熵', '任务目标优化', '-'],
        ['L_mask', 'BCE + Dice', '掩码空间质量', '1.0 / 2.0'],
        ['L_decouple', 'cos^2(C_F, C_B)', '前景/背景方向正交', '0.05 / 0.1'],
        ['L_budget', 'max(r_min-r,0)^2 + max(r-r_max,0)^2', '掩码比例约束[0.1,0.7]', '0.01 / 0.02'],
    ]
)

doc.add_paragraph('伪GT掩码生成：选择分数高于"均值+标准差"的patch作为前景。')

doc.add_page_break()

# ==================== 五、实验设置 ====================
doc.add_heading('五、实验设置', level=1)

doc.add_heading('5.1 数据集', level=2)
add_table(doc,
    ['数据集', '训练集', '验证集', '类别数', '用途'],
    [
        ['PASCAL VOC 2012', '5,717张', '1,449张', '20', '主实验+消融+压力测试'],
        ['COCO 2017', '118,287张', '5,000张', '80', '泛化性验证'],
    ]
)

doc.add_heading('5.2 骨干网络', level=2)
add_table(doc,
    ['骨干', '参数量', '预训练数据', '隐藏维度D', '用途'],
    [
        ['DINO ViT-S/16', '~22M', 'ImageNet', '384', '物体发现'],
        ['OpenCLIP ViT-B/16', '~86M', 'LAION2B', '768', '开放词汇分割'],
    ]
)

doc.add_heading('5.3 训练配置', level=2)
add_table(doc,
    ['超参数', 'DINO ViT-S/16', 'OpenCLIP ViT-B/16'],
    [
        ['提取的层', '[3, 6, 9, 12]', '[3, 6, 9, 12]'],
        ['图像尺寸', '224x224', '224x224'],
        ['patch大小', '16x16', '16x16'],
        ['隐藏维度', '256', '256'],
        ['batch_size', '16', '16'],
        ['训练轮数', '50', '50'],
        ['学习率(VOC)', '1e-4', '5e-5'],
        ['学习率(COCO)', '5e-5', '-'],
        ['权重衰减', '1e-4', '1e-4'],
        ['学习率调度', '余弦退火', '余弦退火'],
        ['梯度裁剪', 'max_norm=1.0', 'max_norm=1.0'],
        ['混合精度', '开启(VOC)/关闭(COCO)', '开启'],
        ['lambda_fg', '1.0(VOC)/0.05(COCO)', '2.0'],
        ['lambda_decouple', '0.05', '0.1'],
        ['lambda_budget', '0.01', '0.02'],
        ['r_min / r_max', '0.1 / 0.7', '0.1 / 0.7'],
        ['温度初始值', '0.1', '0.1'],
        ['频率sigma', '0.25', '0.25'],
        ['空间核大小', '3x3', '3x3'],
        ['优化器', 'AdamW', 'AdamW'],
    ]
)

doc.add_heading('5.4 评估指标', level=2)
bp(doc, 'CorLoc：预测前景中心落在GT边界框内的比例，衡量定位精度')
bp(doc, 'Mask IoU：前景掩码与GT分割掩码的交并比（含精确率和召回率），衡量掩码质量')
bp(doc, 'PiB（Point-in-Box）：聚合CLS最相似的patch落在GT框内的比例，衡量语义-空间对齐度')
bp(doc, 'mIoU：开放词汇分割的平均交并比，衡量分割整体质量')

doc.add_page_break()

# ==================== 六、实验结果 ====================
doc.add_heading('六、实验结果', level=1)

doc.add_heading('6.1 主实验：DINO ViT-S/16 + VOC（物体发现）', level=2)
add_table(doc,
    ['方法', 'CorLoc', 'Mask IoU', 'PiB'],
    [
        ['CAM', '62.80%', '3.52%', '21.26%'],
        ['DINO-seg', '68.46%', '8.83%', '24.64%'],
        ['LaSt-ViT', '65.70%', '13.31%', '48.10%'],
        ['DAFB-CLS', '79.43%', '31.27%', '45.07%'],
    ]
)
doc.add_paragraph('DAFB-CLS超过最强基线LaSt-ViT：CorLoc +13.73pp，Mask IoU +17.96pp。')

doc.add_heading('6.2 主实验：OpenCLIP ViT-B/16 + VOC（开放词汇分割）', level=2)
add_table(doc,
    ['指标', '得分'],
    [
        ['最佳mIoU（第49 epoch）', '61.70%'],
        ['CorLoc', '77.29%（1120/1449）'],
        ['Mask IoU', '33.32%（精确率34.41%，召回率85.42%）'],
        ['PiB', '79.64%（1154/1449）'],
    ]
)

doc.add_heading('6.3 主实验：DINO ViT-S/16 + COCO（物体发现）', level=2)
add_table(doc,
    ['方法', 'CorLoc', 'Mask IoU', 'PiB'],
    [
        ['CAM', '59.81%', '4.30%', '24.96%'],
        ['DINO-seg', '67.23%', '9.24%', '29.00%'],
        ['LaSt-ViT', '61.53%', '12.95%', '45.09%'],
        ['DAFB-CLS', '72.96%', '28.73%', '35.44%'],
    ]
)
doc.add_paragraph('DAFB-CLS超过最强基线LaSt-ViT：CorLoc +11.43pp，Mask IoU +15.78pp。')

doc.add_heading('6.4 跨数据集一致性', level=2)
add_table(doc,
    ['提升幅度', 'VOC（20类）', 'COCO（80类）', '一致性'],
    [
        ['CorLoc', '+13.73pp', '+11.43pp', '两个数据集均显著提升'],
        ['Mask IoU', '+17.96pp', '+15.78pp', '两个数据集均大幅提升'],
        ['PiB', '-3.03pp', '-9.65pp', '两个数据集均略有下降'],
    ]
)

doc.add_page_break()

# ==================== 七、消融实验 ====================
doc.add_heading('七、消融实验', level=1)

doc.add_heading('7.1 DINO ViT-S/16 + VOC 消融', level=2)
add_table(doc,
    ['变体', 'CorLoc%', 'MaskIoU%', 'PiB%', '去除的组件'],
    [
        ['baseline_dino', '29.33', '30.76', '42.72', '整个DAFB框架'],
        ['full', '79.43', '31.27', '45.07', '无（完整模型）'],
        ['no_budget', '79.78', '13.91', '53.62', '自适应预算'],
        ['no_cues', '47.48', '30.16', '31.68', '所有4条线索'],
        ['no_dual_cls', '78.40', '31.25', '33.13', '双CLS（单流替代）'],
        ['shared_depth', '77.43', '31.18', '28.36', '独立深度注意力'],
        ['hard_topk', '81.71', '22.68', '51.76', '软掩码'],
    ]
)

doc.add_paragraph('关键发现：')
bp(doc, '多线索融合是基石：去除后CorLoc从79.43%暴跌至47.48%（-31.95pp）')
bp(doc, '自适应预算防止掩码坍缩：去除后Mask IoU从31.27%暴跌至13.91%（-17.36pp）')
bp(doc, '双CLS解耦提升空间定位：去除后PiB从45.07%降至33.13%（-11.94pp）')
bp(doc, '独立深度注意力有帮助：共享后PiB从45.07%降至28.36%（-16.71pp）')
bp(doc, '软掩码优于硬Top-K：硬Top-K的CorLoc更高（81.71%）但Mask IoU远低（22.68% vs 31.27%）')

doc.add_heading('7.2 OpenCLIP ViT-B/16 + VOC 消融', level=2)
add_table(doc,
    ['变体', 'mIoU%', '与基线差异'],
    [
        ['baseline_clip', '59.02', '基线'],
        ['no_budget', '64.80', '+5.78'],
        ['no_cues', '62.61', '+3.59'],
        ['shared_depth', '62.34', '+3.32'],
        ['full', '62.27', '+3.25'],
        ['hard_topk', '61.66', '+2.64'],
        ['no_dual_cls', '61.65', '+2.63'],
    ]
)
doc.add_paragraph('所有DAFB变体均超过基线59.02%，验证框架对分割任务的通用增益。')

doc.add_page_break()

# ==================== 八、压力测试 ====================
doc.add_heading('八、压力测试与失败模式分析', level=1)

doc.add_heading('8.1 四个子集的压力测试', level=2)
add_table(doc,
    ['子集', '样本数', 'FG IoU', 'PiB', '预测FG比例', 'Pearson r'],
    [
        ['稳定背景', '551', '13.02%', '20.87%', '0.484', '0.394'],
        ['纹理前景', '898', '34.42%', '43.88%', '0.586', '0.635'],
        ['小目标', '450', '26.92%', '34.67%', '0.550', '0.645'],
        ['多目标', '436', '36.19%', '43.12%', '0.600', '0.605'],
    ]
)

doc.add_heading('8.2 失败模式分析', level=2)
doc.add_paragraph(
    '稳定背景是主要失败模式：频率稳定性线索将平滑背景（天空、墙面、水面）误判为前景。'
    '预测前景比例顽固维持在0.48-0.60，无论GT真实比例如何。'
    'Pearson相关系数仅0.394（其他子集>0.6），说明模型对这类图像的预测几乎失控。'
)

doc.add_page_break()

# ==================== 九、改进尝试 ====================
doc.add_heading('九、改进尝试与未来方向', level=1)

doc.add_heading('9.1 已尝试的改进（6轮）', level=2)
add_table(doc,
    ['轮次', '策略', '稳定BG FG IoU', '预测FG比例', '结论'],
    [
        ['1', '纹理互补线索', '12.62%', '0.485', '高频残差在平滑区域几乎为零'],
        ['2', '+ r_min=0, lambda_fg=5.0', '12.68%', '0.482', '模型仍预测~50%前景'],
        ['3', '+ 温度初始化0.5', '12.82%', '0.512', '预测比例反而更接近0.5'],
        ['4', '+ GT比例对齐损失', '12.75%', '0.511', '比例损失被其他损失淹没'],
        ['5', '+ tau教师强制', '10.58%', '0.902', 'tau坍缩，比例飙至0.9'],
        ['6', '+ tau预测器对齐', '5.99%', '0.365', '进一步破坏tau预测能力'],
    ]
)
doc.add_paragraph('结论：简单的损失/线索调优无法解决此问题，需要架构层面的根本性改进。')

doc.add_heading('9.2 未来改进方向', level=2)
bp(doc, '增加边缘/梯度线索：用Sobel算子或Canny边缘检测区分平滑前景（有边缘）和平滑背景（无边缘）')
bp(doc, '对比/排序损失：替代阈值掩码，学习patch之间的前景排序关系')
bp(doc, '显式背景先验：使用预训练分割模型（如SAM）的输出作为背景监督')
bp(doc, '扩展骨干网络：测试DINO ViT-B/16和OpenCLIP ViT-L/14')
bp(doc, '补充对比方法：实现TokenCut、MaskCLIP等更多基线')
bp(doc, '效率分析：报告参数量、FLOPs、推理时间')

doc.add_page_break()

# ==================== 十、代码架构 ====================
doc.add_heading('十、代码架构说明', level=1)

doc.add_heading('10.1 目录结构', level=2)
bp(doc, 'dafb_cls/models/ —— 模型定义（dafb_cls_model.py, cues.py, foregroundness.py, adaptive_budget.py, dual_cls.py, depth_attention.py, fusion.py, heads.py, feature_extractor.py）')
bp(doc, 'dafb_cls/losses/ —— 损失函数（mask_losses.py, decouple_loss.py, budget_loss.py）')
bp(doc, 'dafb_cls/datasets/ —— 数据集加载（voc.py, coco.py）')
bp(doc, 'dafb_cls/tools/ —— 训练和评估脚本（train_posthoc.py, eval_corloc.py, eval_mask_iou.py, eval_pib.py, eval_lastvit.py, eval_dinoseg.py, eval_cam.py, stress_test.py, run_ablation.py, visualize_masks.py, visualize_depth_attention.py, extract_features.py）')
bp(doc, 'dafb_cls/configs/ —— 配置文件（dino_vits16_voc.yaml, openclip_vitb16_voc.yaml, dino_vits16_coco.yaml, openclip_vitb16_coco.yaml, ablation yaml文件）')
bp(doc, 'checkpoints/ —— 训练好的模型权重')
bp(doc, 'scripts/ —— 数据集下载脚本（download_voc2012.py, download_coco2017.py）')

doc.add_heading('10.2 核心模块说明', level=2)

bold_para(doc, 'MultiLayerFeatureExtractor（feature_extractor.py）')
doc.add_paragraph(
    '通过前向钩子提取骨干网络指定层的中间特征。支持DINO、OpenCLIP、DeiT等骨干。'
    '所有骨干参数冻结（requires_grad_(False)）。'
)

bold_para(doc, 'DualCLSAggregator（dual_cls.py）')
doc.add_paragraph(
    '核心编排模块。计算前景/背景掩码，执行双CLS加权聚合。'
    '包含FrequencyStabilityCue、DepthConsistencyCue、SemanticAlignmentCue、SpatialCompactnessCue、'
    'ForegroundnessHead、AdaptiveBudgetModule、BackgroundnessHead。'
)

bold_para(doc, 'ForegroundBackgroundDepthAttention（depth_attention.py）')
doc.add_paragraph(
    '前景/背景各自独立的深度注意力。包含RMSNorm和DepthAttentionBlock。'
)

bold_para(doc, 'TaskAdaptiveFusionHead（fusion.py）')
doc.add_paragraph('门控融合网络，自适应平衡前景和背景CLS token。')

doc.add_heading('10.3 训练命令', level=2)
doc.add_paragraph('VOC + DINO：')
doc.add_paragraph('python -m dafb_cls.tools.train_posthoc --config dafb_cls/configs/dino_vits16_voc.yaml', style='List Bullet')
doc.add_paragraph('COCO + DINO：')
doc.add_paragraph('set HF_HUB_OFFLINE=1 && python -m dafb_cls.tools.train_posthoc --config dafb_cls/configs/dino_vits16_coco.yaml', style='List Bullet')
doc.add_paragraph('评估命令：')
doc.add_paragraph('python -m dafb_cls.tools.eval_corloc --config <config> --checkpoint <ckpt>', style='List Bullet')
doc.add_paragraph('python -m dafb_cls.tools.eval_mask_iou --config <config> --checkpoint <ckpt>', style='List Bullet')
doc.add_paragraph('python -m dafb_cls.tools.eval_pib --config <config> --checkpoint <ckpt>', style='List Bullet')

doc.add_page_break()

# ==================== 十一、参考文献 ====================
doc.add_heading('十一、完整参考文献', level=1)

refs = [
    '[1] Dosovitskiy, A., et al. "An image is worth 16x16 words: Transformers for image recognition at scale." ICLR 2021.',
    '[2] Touvron, H., et al. "Training data-efficient image transformers & distillation through attention." ICML 2021.',
    '[3] Caron, M., et al. "Emerging properties in self-supervised vision transformers." ICCV 2021.',
    '[4] Oquab, M., et al. "DINOv2: Learning robust visual features without supervision." TMLR 2024.',
    '[5] Radford, A., et al. "Learning transferable visual models from natural language supervision." ICML 2021.',
    '[6] Cherti, M., et al. "Reproducible scaling laws for contrastive language-image learning." CVPR 2023.',
    '[7] Zhou, B., et al. "Learning deep features for discriminative localization." CVPR 2016.',
    '[8] Selvaraju, R.R., et al. "Grad-CAM: Visual explanations from deep networks." ICCV 2017.',
    '[9] Abnar, S. and Zuidema, W. "Quantifying attention flow in transformers." ACL 2020.',
    '[10] Caron, M., et al. "Emerging properties in self-supervised vision transformers." ICCV 2021. (DINO-seg)',
    '[11] Bao, H., et al. "BEiT: BERT pre-training of image transformers." ICLR 2022.',
    '[12] He, K., et al. "Masked autoencoders are scalable vision learners." CVPR 2022.',
    '[13] Chen, X., et al. "An empirical study of training self-supervised vision transformers." ICCV 2021.',
    '[14] Zhou, J., et al. "iBOT: Image BERT pre-training with online tokenizer." ICLR 2022.',
    '[15] Darcet, T., et al. "Vision transformers need registers." ICLR 2024.',
    '[16] Darcet, T., et al. "Vision transformers need more than registers." arXiv 2024.',
    '[17] "Attention Residuals." (论文二)',
    '[18] Zhou, H., et al. "DeepViT: Towards Deeper Vision Transformer." arXiv 2021.',
    '[19] Simeoni, O., et al. "Localizing objects with self-supervised transformers and no labels." BMVC 2021.',
    '[20] Wang, Y., et al. "Cut and learn for unsupervised object detection." CVPR 2023.',
    '[21] Wang, X., et al. "FreeSOLO: Learning to segment objects without annotations." CVPR 2022.',
    '[22] Shi, et al. "LaSt-ViT: Lazy Aggregation via Frequency Stability for Vision Transformers." 2026.',
    '[23] Hung, W.-C., et al. "ScOPS: Self-supervised co-part segmentation." CVPR 2019.',
    '[24] Zhou, C., et al. "Extracting free dense masks from image tokenizers." ICLR 2022.',
    '[25] Liang, F., et al. "Open-vocabulary semantic segmentation with mask-adapted CLIP." CVPR 2023.',
    '[26] Ghiasi, G., et al. "Scaling open-vocabulary image segmentation." ECCV 2022.',
    '[27] Xu, J., et al. "GroupViT: Semantic segmentation emerges from text supervision." CVPR 2022.',
    '[28] Choe, J. and Shim, H. "Attention-based dropout layer for WSOL." CVPR 2019.',
    '[29] Wang, H., et al. "Score-CAM: Score-weighted visual explanations." CVPR 2020.',
    '[30] Li, K., et al. "Tell me where to look: Guided attention inference network." CVPR 2018.',
    '[31] Zhu, X., et al. "Deformable DETR." ICLR 2021.',
    '[32] Vaswani, A., et al. "Attention is all you need." NeurIPS 2017.',
    '[33] Katharopoulos, A., et al. "Transformers are RNNs: Fast autoregressive transformers with linear attention." ICML 2020.',
    '[34] Rao, Y., et al. "Global filter networks for image classification." NeurIPS 2021.',
    '[35] Mezic, I. "Spectral properties of dynamical systems." Nonlinear Dynamics 2005.',
    '[36] He, K., et al. "Deep residual learning for image recognition." CVPR 2016.',
]
for ref in refs:
    doc.add_paragraph(ref)

# Save
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'DAFB_CLS_完整项目文档.docx')
doc.save(output_path)
print(f'已保存: {output_path}')
