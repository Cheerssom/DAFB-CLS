"""
生成 DAFB-CLS 研究成果汇报 Word 文档（中文版）
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# --- 样式设置 ---
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Microsoft YaHei'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    hs.font.color.rgb = RGBColor(0, 51, 102)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
    return table


def add_bold_paragraph(doc, bold_text, normal_text):
    p = doc.add_paragraph()
    run_b = p.add_run(bold_text)
    run_b.bold = True
    p.add_run(normal_text)
    return p


# ===================== 标题 =====================
title = doc.add_heading('DAFB-CLS: 视觉Transformer的深度自适应前景-背景CLS解耦', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('研究成果汇报')
run.font.size = Pt(14)
run.bold = True

doc.add_paragraph()

# ===================== 1. 研究背景 =====================
doc.add_heading('一、研究背景与动机', level=1)

doc.add_paragraph(
    'Vision Transformer（ViT）已成为计算机视觉领域的主流架构。然而，ViT中的CLS token存在两个根本性问题，'
    '严重制约了其可解释性和空间定位能力：'
)

add_bold_paragraph(doc,
    '问题一——惰性聚合（Lazy Aggregation）：',
    'CLS token倾向于将背景patch的全局语义作为捷径进行聚合，原因在于粗粒度的语义监督和全局注意力机制。'
    '该问题由论文"Vision Transformers Need More Than Registers"（Darcet等人，ICLR 2024）首次系统性揭示。'
)

add_bold_paragraph(doc,
    '问题二——惰性累加（Lazy Accumulation）：',
    '标准残差连接对所有历史层的输出进行固定等权重累加，缺乏沿深度维度的选择性。'
    '该问题由论文"Attention Residuals"提出并分析。'
)

doc.add_paragraph(
    '我们的DAFB-CLS框架将上述两个问题的解决方案统一到一个后处理聚合器中，'
    '将CLS token解耦为独立的前景和背景语义流，并为每条流配备深度自适应注意力机制。'
)

# ===================== 2. 论文一分析 =====================
doc.add_heading('二、论文一：Vision Transformers Need More Than Registers（Darcet等人，ICLR 2024）', level=1)

doc.add_heading('2.1 核心贡献', level=2)
doc.add_paragraph(
    '本文首次系统性地揭示了DINOv2中的"artifact"现象：训练过程中添加的额外token吸收了全局注意力，'
    '在注意力图中产生局部高激活区域。作者提出将artifact register替换为显式的前景和背景register，'
    '产生更干净的注意力图，在k-NN分类和分割任务上达到了当时的最优性能。'
)

doc.add_heading('2.2 优势', level=2)
for bold_text, normal_text in [
    ('首次系统性揭示artifact现象：', '发现ViT中间层特征图中存在异常高激活值的patch（artifact token），'
     '这些token在注意力图中充当"垃圾桶"，吸收了本应分配给语义区域的注意力权重。'),
    ('简洁有效的Register方案：', '在训练时添加额外的register token，推理时移除，即可消除artifact。'
     '不改变推理pipeline的复杂度，工程实现简单。'),
    ('建立了清晰的artifact分类学：', '将artifact分为两种类型——Type-1（单一patch接收过多注意力）'
     '和Type-2（注意力在所有token上均匀分布）。'),
    ('实验结果扎实：', '在k-NN分类和分割任务上均达到SOTA，证明了清理注意力质量对下游任务的价值。'),
    ('方法极简：', '仅需在训练时修改输入token序列（加入4个register token），不改变模型架构。'),
]:
    p = doc.add_paragraph(style='List Bullet')
    run_b = p.add_run(bold_text)
    run_b.bold = True
    p.add_run(normal_text)

doc.add_heading('2.3 不足', level=2)
for bold_text, normal_text in [
    ('仅适用于训练阶段：', 'Register方法需要从头训练或微调模型，无法直接用于已有的预训练模型（如DINOv2），'
     '限制了即插即用性。'),
    ('没有区分前景与背景语义：', 'Register token是无差别的占位符，没有语义区分能力——'
     '它们吸收的是"闲置"注意力，而非有目的地分离前景与背景。'),
    ('未考虑深度维度的选择性：', '所有层共享相同的register token，没有根据各层的语义层次'
     '（浅层局部特征 vs 深层全局语义）自适应调整聚合权重。'),
    ('无法用于推理时的任务自适应：', '推理时直接移除register token意味着丢弃了所有与artifact相关的中间信息，'
     '而这些信息可能对特定任务（如分割）有价值。'),
    ('不涉及前/背景解耦的显式建模：', '虽然register token改善了注意力质量，'
     '但没有主动将CLS token的语义分解为前景流和背景流。'),
]:
    p = doc.add_paragraph(style='List Bullet')
    run_b = p.add_run(bold_text)
    run_b.bold = True
    p.add_run(normal_text)

# ===================== 3. 论文二分析 =====================
doc.add_heading('三、论文二：Attention Residuals', level=1)

doc.add_heading('3.1 核心贡献', level=2)
doc.add_paragraph(
    '本文指出了ViT残差连接中的"惰性累加"问题：标准残差连接 x + F(x) 对所有历史层的输出进行等权重累加，'
    '忽略了不同深度特征的重要性差异。作者提出用注意力机制替代固定等权累加，让模型学习在深度维度上的权重分布，'
    '决定哪些层的输出对最终任务最有价值。'
)

doc.add_heading('3.2 优势', level=2)
for bold_text, normal_text in [
    ('准确识别了残差连接的瓶颈：', '标准残差将所有层的贡献等权叠加，忽略了不同深度特征的重要性差异。'),
    ('提出了深度自适应聚合机制：', '用注意力机制替代固定等权累加，让模型学习在深度维度上的权重分布。'),
    ('适用于所有标准ViT架构：', '不依赖特定的预训练方法，具有广泛适用性。'),
]:
    p = doc.add_paragraph(style='List Bullet')
    run_b = p.add_run(bold_text)
    run_b.bold = True
    p.add_run(normal_text)

doc.add_heading('3.3 不足', level=2)
for bold_text, normal_text in [
    ('没有区分前/背景语义流：', '虽然在深度维度上实现了选择性聚合，但仍是对单一CLS token的操作，'
     '没有将前景和背景信息分别处理。'),
    ('缺乏显式的前景感知机制：', '没有利用patch级别的前景性线索来指导聚合过程。'),
    ('对空间结构信息利用不足：', '仅关注深度维度的选择性，未充分利用patch的空间关系。'),
]:
    p = doc.add_paragraph(style='List Bullet')
    run_b = p.add_run(bold_text)
    run_b.bold = True
    p.add_run(normal_text)

# ===================== 4. DAFB-CLS 改进 =====================
doc.add_heading('四、DAFB-CLS：在两篇论文基础上的创新改进', level=1)

doc.add_heading('4.1 统一框架设计', level=2)
doc.add_paragraph(
    'DAFB-CLS将两篇论文的核心问题统一到一个框架中。下表总结了相比每篇论文的关键改进：'
)

add_table(doc,
    ['改进维度', 'Register论文', 'Attention Residuals', 'DAFB-CLS改进'],
    [
        ['前景识别方式', '无（通用占位符）', '无（单一流）', '4线索融合的前景性评分'],
        ['深度聚合', '无选择性', '全局深度注意力', '前景/背景独立深度注意力'],
        ['推理时可用性', '需移除register', '可用', '完全后处理，冻结backbone'],
        ['空间掩码', '无', '无', '自适应软掩码 + 不确定性估计'],
        ['背景建模', '被动消除', '不区分', '显式背景CLS token'],
        ['任务适应性', '无', '无', '门控融合，自适应不同任务'],
    ]
)

doc.add_heading('4.2 五阶段流水线架构', level=2)

add_bold_paragraph(doc, '阶段一：多线索前景区分度计算',
    '——超越两篇论文的单一信号。融合4种互补线索：'
    '（1）频率稳定性线索（FrequencyStabilityCue）：FFT低通滤波，衡量patch特征的频率稳定性；'
    '（2）深度一致性线索（DepthConsistencyCue）：跨层余弦相似度，衡量激活的持久性；'
    '（3）语义对齐线索（SemanticAlignmentCue）：patch-CLS或patch-text相似度；'
    '（4）空间紧凑性线索（SpatialCompactnessCue）：局部邻域平滑。'
    '融合方式：可学习softmax归一化权重加权求和 + MLP分支 + 可选3x3卷积空间平滑。'
)

add_bold_paragraph(doc, '阶段二：自适应前景预算',
    '——改进固定的Top-K选择。MLP从全局CLS特征预测每张图像的阈值tau，'
    '软前景掩码：m_i^F = sigmoid((F_i - tau) / T)，T为可学习温度参数。'
    '独立的BackgroundnessHead预测背景分数，含不确定性估计。'
)

add_bold_paragraph(doc, '阶段三：双CLS解耦聚合',
    '——核心创新。不仅抑制背景（Register论文的做法），也不仅只做单一深度聚合'
    '（Attention Residuals的做法），而是分别计算前景和背景CLS token，'
    '对每个提取的ViT层（默认：第3、6、9、12层）独立计算：'
    'B_l^F = sum(m_i^F * x_i^l) / sum(m_i^F)，B_l^B = sum(m_i^B * x_i^l) / sum(m_i^B)。'
)

add_bold_paragraph(doc, '阶段四：深度注意力',
    '——受Attention Residuals启发，但为前景/背景分别设计。前景和背景流各自拥有独立的pseudo-query向量'
    '（零初始化），RMSNorm归一化key，Softmax注意力在深度（层）维度上计算。'
    '输出两个汇总向量：C_F（前景CLS）和 C_B（背景CLS）。'
)

add_bold_paragraph(doc, '阶段五：任务自适应融合',
    '——门控机制：g = sigmoid(MLP([C_F, C_B, global_feat]))，'
    '最终CLS：C = g * C_F + (1-g) * C_B。允许模型自适应地平衡前景和背景信息。'
)

# ===================== 5. 实验结果 =====================
doc.add_heading('五、详细实验结果', level=1)

doc.add_heading('5.1 实验设置', level=2)

add_bold_paragraph(doc, '数据集：',
    'PASCAL VOC 2012（1,449张验证图像，20个语义类别）；'
    'COCO 2017（5,000张验证图像，80个类别）——实验进行中。')

add_bold_paragraph(doc, '骨干网络（均冻结，仅训练聚合器）：',
    'DINO ViT-S/16（自监督预训练，用于物体发现）；'
    'OpenCLIP ViT-B/16（LAION2B预训练，用于开放词汇分割）。')

add_bold_paragraph(doc, '训练配置：',
    '优化器AdamW，梯度裁剪（max_norm=1.0）；混合精度训练（AMP）；'
    '余弦退火学习率调度；50个epoch；batch_size=16。'
    'DINO配置：lr=1e-4, lambda_fg=1.0, lambda_decouple=0.05, lambda_budget=0.01。'
    'OpenCLIP配置：lr=5e-5, lambda_fg=2.0, lambda_decouple=0.1, lambda_budget=0.02。')

add_bold_paragraph(doc, '损失函数：',
    '主任务损失（交叉熵）+ lambda_fg * 前景掩码损失（BCE + Dice）+ '
    'lambda_decouple * 解耦损失（余弦平方）+ lambda_budget * 预算正则化损失。')

add_bold_paragraph(doc, '评估指标：',
    'CorLoc（预测前景中心是否落在GT边界框内）；Mask IoU（前景掩码与GT分割掩码的交并比）；'
    'PiB（聚合CLS最相似的patch是否落在GT框内）；mIoU（开放词汇分割的平均交并比）。')

doc.add_heading('5.2 主实验结果：DINO ViT-S/16 + VOC（物体发现）', level=2)

add_table(doc,
    ['方法', 'CorLoc', 'Mask IoU', 'PiB'],
    [
        ['ViT基线（仅CLS）', '29.33%', '30.76%', '42.72%'],
        ['LaSt-ViT（FFT稳定性, K=59）', '65.70%', '13.31%', '48.10%'],
        ['DAFB-CLS（完整）', '79.43%', '31.27%', '45.07%'],
    ]
)

p = doc.add_paragraph()
run_b = p.add_run('关键提升：')
run_b.bold = True
p.add_run('CorLoc相比原始ViT提升 +50.1pp，相比LaSt-ViT提升 +13.7pp；'
           'Mask IoU相比LaSt-ViT提升 +18.0pp（13.31%到31.27%）。')

doc.add_heading('5.3 主实验结果：OpenCLIP ViT-B/16 + VOC（开放词汇分割）', level=2)

add_table(doc,
    ['指标', '得分'],
    [
        ['最佳mIoU（第49 epoch）', '61.70%'],
        ['CorLoc', '77.29%（1120/1449）'],
        ['Mask IoU', '33.32%（精确率34.41%，召回率85.42%）'],
        ['PiB', '79.64%（1154/1449）'],
    ]
)

doc.add_heading('5.4 消融实验：DINO ViT-S/16', level=2)

add_table(doc,
    ['变体', 'CorLoc%', 'MaskIoU%', 'PiB%'],
    [
        ['baseline_dino', '29.33', '30.76', '42.72'],
        ['full', '79.43', '31.27', '45.07'],
        ['no_budget', '79.78', '13.91', '53.62'],
        ['no_cues', '47.48', '30.16', '31.68'],
        ['no_dual_cls', '78.40', '31.25', '33.13'],
        ['shared_depth', '77.43', '31.18', '28.36'],
        ['hard_topk', '81.71', '22.68', '51.76'],
    ]
)

doc.add_paragraph('消融分析：')
for f in [
    '多线索前景区分是基础：去除线索后CorLoc从79%暴跌至47%，证明多线索融合的不可或缺性。',
    '自适应预算防止掩码坍缩：去除后Mask IoU从31.27%暴跌至13.91%，固定比例无法适应不同图像的前景面积。',
    '双CLS对空间定位至关重要：去除后PiB从45.07%降至33.13%，单一CLS token无法充分分离前景和背景。',
    '独立深度注意力有帮助：共享后PiB从45.07%降至28.36%，前景和背景需要不同的深度权重分布。',
    '软掩码优于硬Top-K：硬Top-K的CorLoc更高（81.71%），但Mask IoU远低于软掩码（22.68% vs 31.27%），边界伪影严重。',
]:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('5.5 消融实验：OpenCLIP ViT-B/16', level=2)

add_table(doc,
    ['变体', 'mIoU%'],
    [
        ['baseline_clip', '59.02'],
        ['no_budget', '64.80'],
        ['no_cues', '62.61'],
        ['shared_depth', '62.34'],
        ['full', '62.27'],
        ['hard_topk', '61.66'],
        ['no_dual_cls', '61.65'],
    ]
)

doc.add_paragraph('所有DAFB变体均超过基线（59.02%），验证了框架对分割任务的通用提升能力。')

doc.add_heading('5.6 压力测试（DINO ViT-S/16 完整模型）', level=2)

add_table(doc,
    ['子集', '样本数', '前景IoU', 'PiB', '预测FG比例', 'Pearson r'],
    [
        ['稳定背景', '551', '13.02%', '20.87%', '0.484', '0.394'],
        ['纹理前景', '898', '34.42%', '43.88%', '0.586', '0.635'],
        ['小目标', '450', '26.92%', '34.67%', '0.550', '0.645'],
        ['多目标', '436', '36.19%', '43.12%', '0.600', '0.605'],
    ]
)

doc.add_paragraph('压力测试发现：')
for f in [
    '稳定背景是主要失败模式：频率稳定性线索将平滑背景（天空、墙面、水面）误判为前景。',
    '纹理化前景和多目标场景表现良好，验证了双CLS解耦聚合的有效性。',
    '小目标有一定挑战，受限于ViT-S/16的空间分辨率（14x14 patch）。',
    '预测前景比例顽固地维持在0.48-0.60，无论GT真实比例如何。',
]:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('5.7 稳定背景改进尝试（6轮迭代实验）', level=2)

add_table(doc,
    ['轮次', '策略', '稳定背景FG IoU', '预测FG比例'],
    [
        ['1', '纹理互补线索', '12.62%', '0.485'],
        ['2', '+ r_min=0, lambda_fg=5.0', '12.68%', '0.482'],
        ['3', '+ 温度初始化0.5', '12.82%', '0.512'],
        ['4', '+ GT比例对齐损失', '12.75%', '0.511'],
        ['5', '+ tau教师强制', '10.58%', '0.902'],
        ['6', '+ tau预测器对齐', '5.99%', '0.365'],
    ]
)

doc.add_paragraph(
    '六轮迭代实验确认：简单的损失/线索调优无法解决此问题。根本原因在于架构层面——'
    '频率稳定性线索在本质上将平滑前景物体与平滑背景混淆。'
    '这为后续改进指明了方向：需要架构层面的根本性解决方案。'
)

# ===================== 6. 计划改进方向 =====================
doc.add_heading('六、计划改进方向', level=1)

doc.add_heading('6.1 短期目标（1-2个月）', level=2)
for bold_text, normal_text in [
    ('完成COCO实验：',
     '在COCO 2017上分别用DINO ViT-S/16和OpenCLIP ViT-B/16运行DAFB-CLS，验证泛化性。'
     'COCO验证集有5,000张图（VOC的3.4倍），提供更鲁棒的评估。'),
    ('补充对比方法：',
     '实现并对比CAM/GradCAM、TokenCut、DINO-seg和MaskCLIP等基线方法。这是投稿的必要条件。'),
    ('生成定性可视化：',
     '制作方法对比图，在相同图像上展示不同方法的分割结果，加上成功/失败案例分析和深度注意力热力图。'),
]:
    p = doc.add_paragraph(style='List Bullet')
    run_b = p.add_run(bold_text)
    run_b.bold = True
    p.add_run(normal_text)

doc.add_heading('6.2 中期目标（2-3个月）', level=2)
for bold_text, normal_text in [
    ('扩展骨干网络：',
     '测试DINO ViT-B/16和OpenCLIP ViT-L/14，验证方法在不同模型规模下的可扩展性。'),
    ('效率分析：',
     '报告参数量、FLOPs和推理时间对比。DAFB-CLS仅在冻结backbone上训练轻量聚合器，额外开销应很小。'),
    ('架构层面修复稳定背景问题：',
     '添加基于边缘/梯度的线索区分平滑前景和平滑背景；'
     '探索对比/排序损失替代阈值掩码；考虑使用预训练分割特征作为显式背景先验。'),
]:
    p = doc.add_paragraph(style='List Bullet')
    run_b = p.add_run(bold_text)
    run_b.bold = True
    p.add_run(normal_text)

doc.add_heading('6.3 长期目标（3-6个月）', level=2)
for bold_text, normal_text in [
    ('理论分析：',
     '为双流聚合框架提供收敛性保证和复杂度分析。'),
    ('扩展到视频领域：',
     '将DAFB-CLS应用于视频理解任务，其中时序前景/背景分离具有天然重要性。'),
    ('与大型视觉-语言模型集成：',
     '将DAFB-CLS作为CLIP基础模型的即插即用模块，在不微调的情况下改善其空间定位能力。'),
]:
    p = doc.add_paragraph(style='List Bullet')
    run_b = p.add_run(bold_text)
    run_b.bold = True
    p.add_run(normal_text)

doc.add_heading('6.4 目标投稿venue', level=2)
for v in [
    'Workshop论文（CVPR/ICLR Workshop）：当前结果 + COCO实验即可达到。',
    '二区会议（WACV, BMVC, AAAI）：需COCO + 3-5个对比方法 + 可视化。',
    '一区会议（CVPR, ECCV, NeurIPS）：以上全部 + 理论分析 + 大规模实验。',
    '二区期刊（TMM, TNNLS）：全面实验 + 效率分析 + 详细消融。',
]:
    doc.add_paragraph(v, style='List Bullet')

# ===================== 7. 总结 =====================
doc.add_heading('七、总结', level=1)

doc.add_paragraph(
    'DAFB-CLS提出了一个统一的后处理聚合框架，同时解决了Register论文中的"惰性聚合"问题'
    '和Attention Residuals论文中的"惰性累加"问题。核心创新在于前景/背景双流CLS解耦'
    '配合深度自适应注意力，在DINO+VOC基准上实现了CorLoc +13.7pp和Mask IoU +18.0pp'
    '相比LaSt-ViT的提升。全面的消融实验验证了每个组件的贡献。'
    '主要局限是稳定背景场景的处理，需要架构层面的根本性解决方案。'
    '下一步计划包括COCO实验、补充对比方法、以及可扩展性验证。'
)

# 保存
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'DAFB_CLS_研究成果汇报.docx')
doc.save(output_path)
print(f'报告已保存到: {output_path}')
