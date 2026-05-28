# -*- coding: utf-8 -*-
"""
Generate Related Work + Methods docx for DAFB-CLS paper.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re, os


def set_cell_font(cell, size=10, bold=False, italic=False, name="Times New Roman"):
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.size = Pt(size)
            r.font.bold = bold
            r.font.italic = italic
            r.font.name = name
            r.element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        set_cell_font(cell, size=10, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            set_cell_font(cell, size=10)
            if j > 0:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


def set_paragraph_format(p, space_after=6, space_before=3, first_indent=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = align
    if first_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_indent)


def add_para(doc, text, bold=False, italic=False, size=11, indent=0.75, space_after=6, space_before=3):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = "Times New Roman"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    set_paragraph_format(p, space_after=space_after, space_before=space_before, first_indent=indent)
    return p


def add_title(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = "Times New Roman"
        r.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        r.font.color.rgb = RGBColor(0, 0, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h


def add_section(doc, num, title):
    return add_title(doc, f"{num}  {title}", level=1)


def add_sub(doc, num, title):
    return add_title(doc, f"{num}  {title}", level=2)


def add_subsub(doc, title):
    return add_title(doc, title, level=3)


# ═══════════════════════════════════════════════════════════════
#   RELATED WORK
# ═══════════════════════════════════════════════════════════════

def write_related_work(doc):
    add_section(doc, "2", "Related Work")

    # ── 2.1 ──
    add_sub(doc, "2.1", "Vision Transformers and Dense Feature Representations")

    add_subsub(doc, "2.1.1  From Global Attention to Structured Variants")
    add_para(doc,
        "The Vision Transformer (ViT) [Dosovitskiy et al., 2021] demonstrated that pure attention-based "
        "architectures can match or exceed convolutional networks on image classification when pretrained "
        "at sufficient scale. By treating image patches as tokens and appending a learnable [CLS] token, "
        "ViT leverages global self-attention to build a holistic image representation from which the [CLS] "
        "token aggregates information across all spatial locations. This design has since become the "
        "foundation of numerous vision and multimodal models, including CLIP [Radford et al., 2021], "
        "DINO [Caron et al., 2021], and BEiT [Bao et al., 2022].")
    add_para(doc,
        "However, global quadratic attention also introduces substantial computational overhead, motivating "
        "a line of work on structured sparse attention. PVT [Wang et al., 2021] progressively reduces "
        "spatial resolution across stages via spatial-reduction attention, while Swin Transformer [Liu et al., "
        "2021] restricts attention within local windows with periodic shifted-window cross-region "
        "communication. These architectural innovations significantly improve efficiency and have "
        "demonstrated strong performance on dense prediction tasks, yet they impose a fixed spatial "
        "attention topology that is data-agnostic. DAT [Xia et al., 2022] addressed this limitation through "
        "deformable attention, dynamically predicting sampling offsets from input content to attend to "
        "task-relevant spatial locations. MViT [Fan et al., 2021] introduced multi-scale feature hierarchies "
        "with pooled attention, progressively expanding channel dimensions while reducing spatial resolution. "
        "Despite these advances in architecture design, a separate but equally important class of problems "
        "concerns the quality and interpretability of dense feature representations that ViT produces.")

    add_subsub(doc, "2.1.2  Artifacts and Pathologies in ViT Dense Features")
    add_para(doc,
        "While ViT excels at global image-level tasks, a growing body of evidence reveals systematic "
        "pathologies in its dense (patch-level) feature maps. Darcet et al. [2024] first documented the "
        "\"high-norm token\" phenomenon: in self-supervised ViTs such as DINOv2, certain background tokens "
        "inflate to extreme activation norms, creating spatial artifacts that degrade dense prediction "
        "quality. They proposed adding register tokens—unused positions that absorb background noise—as a "
        "pragmatic fix. While effective at eliminating anomalous high-norm patches, register tokens treat "
        "the symptom without addressing the underlying cause, and do not improve or may even reduce "
        "point-in-box localization accuracy.")
    add_para(doc,
        "Concurrently, label-supervised ViTs exhibit attention deficit: the [CLS] attention map often "
        "disperses across irrelevant background regions rather than focusing on class-relevant objects "
        "[Naseer et al., 2021]. In text-supervised models like CLIP, patch features may not align well with "
        "grounded textual semantics, leading to poor performance on open-vocabulary segmentation and "
        "detection tasks [Park et al., 2022; Zhou et al., 2022]. These seemingly distinct pathologies share "
        "a common root: the coarse-grained supervision signals (image-level labels or image-caption pairs) "
        "do not constrain the spatial specificity of the aggregated representation.")

    add_subsub(doc, "2.1.3  Diagnostic Frameworks for Patch-Level Quality")
    add_para(doc,
        "A critical contribution toward understanding these pathologies was made by Shi et al. [2025], who "
        "proposed two complementary diagnostic metrics. The Patch Score, defined as the cosine similarity "
        "between each patch feature and the global [CLS] representation, quantifies how strongly each "
        "spatial location contributes to the final image-level representation. The Point-in-Box (PiB) "
        "metric evaluates whether the highest-scoring patch falls within the annotated object bounding box. "
        "Using these metrics, Shi et al. demonstrated that vanilla ViT consistently assigns higher Patch "
        "Scores to background patches than foreground ones—achieving PiB values of only 19–48% compared to "
        "60–80% for convolutional networks. Crucially, they showed that this low PiB emerges early in "
        "training and persists throughout optimization, confirming it as a structural pathology rather than "
        "a training artifact.")

    # ── 2.2 ──
    add_sub(doc, "2.2", "Selective CLS Aggregation: From Frequency Stability to Adaptive Selection")

    add_subsub(doc, "2.2.1  LaSt-ViT: Frequency Stability as a Foreground Proxy")
    add_para(doc,
        "Shi et al. [2025] diagnosed the root cause of ViT dense feature artifacts as Lazy Aggregation: "
        "the [CLS] token, under image-level supervision, tends to aggregate semantic information from "
        "whichever patches are most easily accessible—typically the majority background patches—rather than "
        "from the semantically informative foreground patches. This occurs because background patches can "
        "absorb foreground information through global attention, and aggregating from many background "
        "patches is a lower-entropy solution that minimizes training loss.")
    add_para(doc,
        "To address Lazy Aggregation, Shi et al. proposed LaSt-ViT (Lazy Strike), which replaces the "
        "default all-patch aggregation with a frequency-based selection mechanism. For each patch feature, "
        "LaSt-ViT applies a 1D FFT along the channel dimension, passes the result through a Gaussian "
        "low-pass filter, and inverse-transforms back. The rationale is that foreground object patches "
        "exhibit more stable (low-frequency) channel-wise activations, while background patches tend toward "
        "higher-frequency, less coherent activations. A channel-wise Top-K pooling then selects the K most "
        "stable patches for each feature channel, and their features are averaged to form the aggregated "
        "representation. This simple yet effective approach eliminates the high-norm phenomenon, "
        "substantially improves PiB, and delivers consistent gains on zero-shot semantic segmentation "
        "(e.g., +1.3 mIoU on ADE20K with CLIP ViT-B/16) and open-vocabulary detection.")
    add_para(doc,
        "While LaSt-ViT represents an important advance, it has several limitations that motivate our "
        "work. First, the fixed Top-K ratio creates a rigid foreground budget that cannot adapt to "
        "images with varying object sizes or complexity. Second, frequency stability alone may conflate "
        "smooth foreground objects (e.g., a white swan) with smooth background regions (e.g., sky or "
        "walls), leading to the stable-background failure mode that we systematically identify and "
        "address. Third, LaSt-ViT does not explicitly model the depth dimension of ViT's multi-layer "
        "representations, missing the opportunity to selectively aggregate information across layers.")

    add_subsub(doc, "2.2.2  Adaptive and Multi-Cue Aggregation")
    add_para(doc,
        "Several concurrent works have explored alternatives to fixed aggregation strategies. TokenCut "
        "[Wang et al., 2023] leverages self-similarity graphs and normalized cuts to segment foreground "
        "tokens without any training, achieving strong object discovery results. However, graph-based "
        "methods incur non-trivial computational overhead on large token sets. MaskCLIP [Zhou et al., 2022] "
        "modifies CLIP's attention to produce more spatially grounded patch features through text-guided "
        "masking, but requires text supervision at inference time. FreeDA [Zhang et al., 2024] proposes "
        "training-free debiased adaptation for CLIP, using covariance estimation to suppress background "
        "biases, yet operates at the output feature level rather than modifying the aggregation mechanism.")
    add_para(doc,
        "From a complementary direction, several works have explored adaptive token selection for "
        "efficiency. DynamicViT [Rao et al., 2021] and ATS [Fayyaz et al., 2022] learn to prune "
        "unimportant tokens progressively through the network, improving inference speed but not "
        "specifically targeting the CLS aggregation bias. ToMe [Bolya et al., 2023] merges similar tokens "
        "to reduce computation while preserving visual fidelity. These works demonstrate that intelligent "
        "spatial selection improves ViT representations, but they focus on efficiency rather than on "
        "correcting the fundamental CLS aggregation pathology.")
    add_para(doc,
        "Our work builds upon LaSt-ViT's insight of frequency-guided selection while extending it in "
        "three critical dimensions: (i) we enrich the foreground signal with multiple complementary cues "
        "beyond frequency stability alone; (ii) we replace the fixed Top-K with an image-adaptive soft "
        "mask predicted from the global image feature; and (iii) we decouple the foreground and background "
        "streams rather than simply suppressing background, enabling each stream to independently develop "
        "task-relevant semantics.")

    # ── 2.3 ──
    add_sub(doc, "2.3", "Depth-Wise Information Aggregation in Deep Transformers")

    add_subsub(doc, "2.3.1  Residual Connections as Implicit Aggregators")
    add_para(doc,
        "The residual connection, introduced by He et al. [2016] for deep convolutional networks and "
        "adopted universally in Transformer architectures, serves a dual role that is often underappreciated. "
        "Beyond providing an identity path for stable gradient flow—enabling the training of networks "
        "with hundreds or thousands of layers—residual connections implicitly define how information "
        "accumulates across depth. In the PreNorm formulation standard in modern Transformers, the hidden "
        "state at layer l is computed as h_l = h_{l-1} + f_l(h_{l-1}), where f_l is the l-th transformer "
        "block. Unrolling this recursion reveals that the input to any layer is the sum of the initial "
        "embedding and all previous layer outputs, each weighted uniformly by 1.")
    add_para(doc,
        "This uniform accumulation creates a fundamental tension in deep Transformers. As network depth "
        "increases, the hidden state magnitude grows approximately as O(L), causing the relative contribution "
        "of any single layer to be progressively diluted—a phenomenon termed PreNorm dilution [Xiong et al., "
        "2020]. Early-layer features encoding low-level patterns become entangled in a monolithic "
        "representation, and later layers cannot selectively retrieve specific earlier information. This "
        "problem is particularly acute in Vision Transformers, where different layers capture fundamentally "
        "different levels of abstraction: early layers encode textures and edges, middle layers capture "
        "parts and patterns, and later layers encode object-level semantics [Raghu et al., 2021; "
        "Nayman et al., 2022].")

    add_subsub(doc, "2.3.2  Gated and Selective Residual Connections")
    add_para(doc,
        "Recognizing the limitations of fixed residual addition, several works have introduced learnable "
        "gating mechanisms. Highway Networks [Srivastava et al., 2015] pioneered this direction by adding "
        "data-dependent gates that control information flow through layers, predating the Transformer "
        "era. DenseNet [Huang et al., 2017] adopted an alternative strategy of concatenating rather than "
        "adding features from all preceding layers, giving later layers explicit access to all earlier "
        "representations at the cost of linearly growing memory.")
    add_para(doc,
        "In the Transformer context, GPT-2 [Radford et al., 2019] introduced learnable scaling factors "
        "on residual paths, initialized to $1/\\sqrt{N}$ where N is the number of residual connections, "
        "to stabilize training of deep models. More recently, DenseFormer [Petrov et al., 2024] extended "
        "this idea by learning per-layer static weights for all historical representations, providing "
        "content-independent but layer-specific aggregation. While an improvement over uniform addition, "
        "static weights cannot adapt to input-specific information needs—different inputs may require "
        "different layers to contribute at different magnitudes. SubLN [Bachmann et al., 2024] revisited "
        "PreNorm residual scaling and demonstrated that careful normalization within residual branches "
        "can partially mitigate the dilution effect.")

    add_subsub(doc, "2.3.3  Attention Residuals: Content-Dependent Depth Aggregation")
    add_para(doc,
        "The most directly relevant prior work on depth-wise aggregation is Attention Residuals (AttnRes) "
        "[Kimi Team, 2025], which fundamentally reconceptualizes the residual connection as a "
        "depth-dimension attention mechanism. In AttnRes, each layer l receives its input not through "
        "uniform summation but through a learned softmax attention over all historical layer outputs: "
        "$h_l = \\sum_i \\alpha_{i \\to l} \\cdot v_i$, where $v_0$ is the token embedding and $v_i$ is "
        "the output of layer $i$. The attention weights $\\alpha$ are computed using a learned "
        "pseudo-query vector $w_l$ as the query and RMSNorm-transformed historical outputs as keys, "
        "enabling content-dependent depth selection.")
    add_para(doc,
        "AttnRes demonstrated that this selective depth aggregation yields a 1.25× compute advantage "
        "on language modeling benchmarks—achieved at essentially zero additional parameter cost, since "
        "the pseudo-queries are simply d-dimensional vectors. The authors further proposed Block "
        "Attention Residuals (Block AttnRes), which partitions layers into blocks and applies inter-block "
        "attention, reducing storage from O(Ld) to O(Nd) for N blocks while preserving most of the "
        "benefits. Empirical analysis on a 48B-parameter model showed that Block AttnRes eliminates the "
        "monotonic growth of hidden-state magnitude and produces more uniform gradient distributions "
        "across depth.")
    add_para(doc,
        "While AttnRes was developed in the context of language models, its core insight—that depth-wise "
        "aggregation should be content-adaptive rather than fixed—is directly applicable to Vision "
        "Transformers. Indeed, the multi-layer representations of ViT are even more heterogeneous than "
        "those of language models, with each layer capturing qualitatively different visual information. "
        "However, AttnRes addresses all tokens uniformly through a single aggregation mechanism, without "
        "distinguishing between foreground and background information streams. Our method extends this "
        "paradigm by coupling depth-wise attention with foreground-background decoupling, enabling "
        "independent depth aggregation for each semantic stream.")

    # ── 2.4 ──
    add_sub(doc, "2.4", "Foreground-Background Separation in Visual Representation Learning")

    add_subsub(doc, "2.4.1  Class Activation Mapping and CAM-Based Methods")
    add_para(doc,
        "Separating foreground from background is a longstanding challenge in computer vision. Class "
        "Activation Mapping (CAM) [Zhou et al., 2016] and its extensions—GradCAM [Selvaraju et al., 2017], "
        "GradCAM++ [Chattopadhyay et al., 2018], and ScoreCAM [Wang et al., 2020]—generate coarse "
        "localization heatmaps from classification networks by back-propagating or weighting activation "
        "maps. These methods have proven invaluable for weakly supervised localization and segmentation, "
        "but they inherit the same background bias as the underlying classifier. When the classification "
        "network relies on background context for prediction—as ViT often does—the CAM naturally "
        "highlights background regions, achieving only 21–25% PiB in our experiments.")
    add_para(doc,
        "DINO-seg [Oquab et al., 2024] extracts self-attention maps from self-supervised ViTs to produce "
        "unsupervised object segmentations, leveraging the observation that DINO's attention heads "
        "naturally attend to semantically meaningful regions. While more spatially precise than CAM, "
        "DINO-seg's attention maps still suffer from background activations and require post-processing "
        "(e.g., thresholding, PCA-based clustering) that introduces additional hyperparameters. Our "
        "approach can be seen as providing a principled, trainable alternative to these post-hoc "
        "foreground extraction methods.")

    add_subsub(doc, "2.4.2  Multi-Cue Visual Saliency and Objectness")
    add_para(doc,
        "Visual saliency detection has a rich history of leveraging multiple complementary cues. Classic "
        "methods combine color contrast, texture uniqueness, and spatial priors to predict fixation points "
        "[Itti et al., 1998; Hou & Zhang, 2007]. More recently, deep saliency models have incorporated "
        "edge detection, boundary cues, and semantic features [Hou et al., 2019]. The principle that no "
        "single cue is sufficient for robust saliency detection—and that cues should be adaptively "
        "weighted based on context—directly motivates our multi-cue foregroundness approach.")
    add_para(doc,
        "In the ViT feature space specifically, several works have noted that different cues capture "
        "complementary aspects of foregroundness. TokenCut [Wang et al., 2023] uses self-similarity "
        "as a graph-theoretic proxy for objectness. LOST [Siméoni et al., 2021] combines self-similarity "
        "with a seeding strategy that progressively refines the foreground region. FreeSeg [Qin et al., "
        "2022] segments images without training by combining feature similarity with spatial "
        "compactness. These methods demonstrate that combining multiple signals yields more robust "
        "foreground estimates than any single measure, but they operate at the feature level without "
        "modifying the underlying aggregation mechanism. Our work integrates multi-cue reasoning directly "
        "into the CLS aggregation pipeline.")

    # ── 2.5 ──
    add_sub(doc, "2.5", "Connections to Post-Hoc Aggregation and Frozen Backbone Methods")

    add_para(doc,
        "An important design principle that our work shares with several prior methods is operating in a "
        "post-hoc manner on frozen pretrained backbones. This paradigm, exemplified by methods like "
        "MaskCLIP [Zhou et al., 2022], Saliency-guided Open-Vocabulary Segmentation [Liang et al., 2023], "
        "and FreeDA [Zhang et al., 2024], offers several practical advantages: it avoids costly "
        "retraining of large foundation models, preserves the backbone's generalization properties, "
        "and allows the aggregation module to be lightweight and fast to train. Our approach follows this "
        "paradigm, training only approximately 0.59M parameters (2.7% of the DINO ViT-S/16 backbone) "
        "while achieving substantial improvements in dense feature quality.")
    add_para(doc,
        "More broadly, our work fits into the emerging trend of studying and correcting information flow "
        "within pretrained Transformers, rather than simply scaling up model or data size. Just as "
        "Attention Residuals revealed that the default residual connection is suboptimal for information "
        "aggregation across depth, we show that the default CLS pooling is suboptimal for spatial "
        "aggregation. Both findings point to the same overarching lesson: the default aggregation "
        "operations in Transformers—whether across layers or across spatial positions—are overly rigid "
        "and fail to exploit the structured information available in the intermediate representations.")

    # ── 2.6 ──
    add_sub(doc, "2.6", "Summary: Positioning Our Work")

    add_para(doc,
        "Table 1 positions our work relative to the most closely related methods along four key "
        "dimensions: spatial selectivity (whether aggregation is content-adaptive), multi-cue "
        "foregroundness (whether multiple complementary signals are integrated), adaptive budget "
        "(whether the foreground proportion is image-dependent), and depth-wise attention (whether "
        "layer aggregation is content-dependent). LaSt-ViT introduced spatial selectivity through "
        "frequency stability but uses fixed Top-K with a single cue. Attention Residuals introduced "
        "depth-wise attention but operates uniformly across spatial positions. Our DAFB-CLS framework "
        "is the first to unify all four dimensions into a single coherent framework for CLS aggregation "
        "in Vision Transformers.")

    add_table(doc,
        ["Method", "Spatial\nSelectivity", "Multi-Cue\nFG", "Adaptive\nBudget", "Depth-wise\nAttention"],
        [
            ["CAM / GradCAM", "No", "No", "No", "No"],
            ["DINO-seg", "Partial", "No", "No", "No"],
            ["TokenCut", "Yes", "Partial", "No", "No"],
            ["LaSt-ViT", "Yes", "No (freq only)", "No (fixed K)", "No"],
            ["Attention Residuals", "No", "No", "N/A", "Yes"],
            ["DenseFormer", "No", "No", "N/A", "Partial (static)"],
            ["DAFB-CLS (Ours)", "Yes", "Yes (4 cues)", "Yes (learned τ)", "Yes (separate F/B)"],
        ],
        col_widths=[3.5, 2.2, 2.5, 2.5, 2.5])


# ═══════════════════════════════════════════════════════════════
#   METHODS
# ═══════════════════════════════════════════════════════════════

def write_methods(doc):
    add_section(doc, "3", "Method")

    add_sub(doc, "3.1", "Problem Formulation and Design Rationale")
    add_para(doc,
        "We consider a frozen Vision Transformer backbone (e.g., DINO ViT-S/16 or OpenCLIP ViT-B/16) "
        "that processes an image $x \\in \\mathbb{R}^{H \\times W \\times 3}$ into a sequence of patch "
        "tokens $\\{x_i\\}_{i=1}^{N}$ plus a [CLS] token, where $N = (H/P) \\times (W/P)$ for patch "
        "size $P$. The backbone consists of $L$ transformer blocks; we extract intermediate features "
        "from a subset of layers $\\mathcal{L} = \\{l_1, l_2, ..., l_K\\}$ (typically $K=4$ layers "
        "equally spaced, e.g., $\\{3, 6, 9, 12\\}$ for a 12-layer ViT).")
    add_para(doc,
        "Our goal is to train a lightweight post-hoc aggregation module that transforms these multi-layer "
        "features into an improved image representation $C$, without modifying or retraining the frozen "
        "backbone. This design is motivated by two complementary insights from the literature: (i) the "
        "default [CLS] pooling suffers from Lazy Aggregation bias toward background tokens [Shi et al., "
        "2025], and (ii) the standard residual connections accumulate layer outputs with fixed uniform "
        "weights, lacking content-adaptive depth selection [Kimi Team, 2025].")
    add_para(doc,
        "Our key design principle is to address both problems jointly through foreground-background "
        "decoupling: rather than suppressing background information, we maintain two separate semantic "
        "streams—a foreground stream $C_F$ that aggregates foreground-identifying features, and a "
        "background stream $C_B$ that aggregates contextual scene features—and learn to fuse them "
        "adaptively based on the task and input. Figure 1 provides an overview of the complete framework.")

    add_sub(doc, "3.2", "Step 1: Multi-Layer Feature Extraction")

    add_para(doc,
        "Given input image $x$, we first pass it through the frozen ViT backbone $\\Phi$ and extract "
        "intermediate patch features via forward hooks registered at the specified layers $\\mathcal{L}$:")
    add_para(doc,
        "$\\{P^l\\}_{l \\in \\mathcal{L}} = \\text{Extract}(\\Phi(x), \\mathcal{L}), \\quad "
        "P^l \\in \\mathbb{R}^{B \\times N \\times D}$",
        italic=True, size=11)
    add_para(doc,
        "where $P^l$ denotes the patch token features (excluding [CLS]) at layer $l$, $B$ is the batch "
        "size, $N$ is the number of patches, and $D$ is the feature dimension. We stack features across "
        "layers to form a 4D tensor $\\mathbf{P} \\in \\mathbb{R}^{B \\times K \\times N \\times D}$. "
        "The [CLS] token feature from the final extracted layer provides the global image representation "
        "$c_{\\text{cls}} \\in \\mathbb{R}^{B \\times D}$. All features are detached from the backbone "
        "computation graph to ensure that backbone parameters remain frozen during training.")

    add_sub(doc, "3.3", "Step 2: Multi-Cue Foregroundness Scoring")

    add_para(doc,
        "The first major component computes a per-patch foregroundness score that indicates how likely "
        "each patch is to belong to a foreground object. Motivated by the observation that no single "
        "signal reliably distinguishes foreground from background across diverse image types, we design "
        "four complementary cues that capture different aspects of foregroundness. Each cue operates on "
        "the multi-layer patch features $\\mathbf{P}$ and produces a per-patch score "
        "$s_k \\in \\mathbb{R}^{B \\times N}$ for $k \\in \\{1, 2, 3, 4\\}$.")

    add_subsub(doc, "3.3.1  Frequency Stability Cue")
    add_para(doc,
        "Inspired by LaSt-ViT's observation that foreground patches exhibit more coherent channel-wise "
        "activations, we measure the stability of each patch's feature vector under spectral low-pass "
        "filtering. For each patch feature $p_i \\in \\mathbb{R}^D$, we compute its 1D FFT along the "
        "channel dimension, apply a Gaussian low-pass filter with cutoff frequency $\\sigma_f$, and "
        "inverse-transform to obtain the filtered version $\\tilde{p}_i$:")
    add_para(doc,
        "$\\tilde{p}_i = \\text{IFFT}(\\text{FFT}(p_i) \\cdot G(\\sigma_f)), \\quad "
        "G(\\sigma_f) = \\exp\\left(-\\frac{f^2}{2\\sigma_f^2}\\right)$",
        italic=True, size=11)
    add_para(doc,
        "The stability score for patch $i$ is then:")
    add_para(doc,
        "$S_i = \\frac{1}{D} \\sum_{d=1}^{D} \\frac{\\tilde{p}_{i,d}}{|\\tilde{p}_{i,d} - p_{i,d}| + \\epsilon}$",
        italic=True, size=11)
    add_para(doc,
        "Higher values indicate that the patch feature changes minimally under spectral filtering, "
        "suggesting stable, low-frequency semantic content typical of foreground objects. We use "
        "$\\sigma_f = 0.25$ as the default cutoff, following LaSt-ViT. The frequency cue operates "
        "on all extracted layers and the final-layer scores $S_i^{(l_K)}$ are used downstream.")

    add_subsub(doc, "3.3.2  Depth Consistency Cue")
    add_para(doc,
        "While the frequency cue captures within-layer stability, the depth consistency cue captures "
        "stability across layers. Patches corresponding to semantically meaningful objects tend to "
        "maintain more consistent representations across network depth, while background patches "
        "exhibit greater inter-layer variation due to the absence of strong semantic anchoring. We "
        "measure depth consistency as the average cosine similarity between each layer's patch feature "
        "and the cross-layer mean:")
    add_para(doc,
        "$\\bar{p}_i = \\frac{1}{K} \\sum_{k=1}^{K} P_{i}^{l_k}, \\quad "
        "D_i = \\frac{1}{K} \\sum_{k=1}^{K} \\frac{P_{i}^{l_k} \\cdot \\bar{p}_i}"
        "{\\|P_{i}^{l_k}\\| \\cdot \\|\\bar{p}_i\\|}$",
        italic=True, size=11)
    add_para(doc,
        "where $P_i^{l_k}$ denotes the feature of patch $i$ at layer $l_k$. This cue complements "
        "frequency stability by capturing temporal (across-depth) rather than spectral (within-channel) "
        "consistency, and is particularly effective for foreground objects whose semantic identity "
        "emerges progressively through the network.")

    add_subsub(doc, "3.3.3  Semantic Alignment Cue")
    add_para(doc,
        "The semantic alignment cue measures how strongly each patch's feature aligns with the global "
        "image semantics, providing a content-based signal that complements the structural signals from "
        "the frequency and depth cues. We compute the average cosine similarity between each patch "
        "feature and the [CLS] token feature across all extracted layers:")
    add_para(doc,
        "$A_i = \\frac{1}{K} \\sum_{k=1}^{K} \\frac{P_i^{l_k} \\cdot c_{\\text{cls}}}"
        "{\\|P_i^{l_k}\\| \\cdot \\|c_{\\text{cls}}\\|}$",
        italic=True, size=11)
    add_para(doc,
        "For text-supervised backbones (e.g., OpenCLIP), we additionally support text-guided semantic "
        "alignment, where the similarity is computed between each patch and the set of text embeddings "
        "$\\{t_j\\}$, using the maximum similarity: $A_i = \\max_j \\cos(P_i^{l_K}, t_j)$. This "
        "provides a more task-specific foreground signal when text features are available.")

    add_subsub(doc, "3.3.4  Spatial Compactness Cue")
    add_para(doc,
        "Foreground objects in natural images typically occupy spatially contiguous regions. The spatial "
        "compactness cue enforces this prior by applying local neighborhood smoothing to the semantic "
        "alignment scores, encouraging patches within the same spatial neighborhood to have similar "
        "foregroundness. Given the semantic score map $s_{\\text{sem}} \\in \\mathbb{R}^{N}$, we "
        "reshape it to the spatial grid $(h \\times w)$ and apply average pooling with kernel size $k_s$:")
    add_para(doc,
        "$C_i = \\text{AvgPool}(\\text{Reshape}(s_{\\text{sem}}, h, w), k_s)_{[i]}$",
        italic=True, size=11)
    add_para(doc,
        "with $k_s = 3$ as default. This cue does not introduce additional learnable parameters "
        "but provides a useful structural inductive bias that complements the pointwise signals "
        "from the other cues.")

    add_subsub(doc, "3.3.5  Learned Cue Fusion")
    add_para(doc,
        "The four cues are fused through a combination of learnable weighting and an MLP refinement "
        "module. Let $\\mathbf{s}_i = [S_i, D_i, A_i, C_i]^T \\in \\mathbb{R}^4$ denote the cue "
        "vector for patch $i$. We first compute a weighted combination with learned softmax-normalized "
        "weights:")
    add_para(doc,
        "$F_i^{\\text{linear}} = \\sum_{k=1}^{4} w_k \\cdot s_{i,k}, \\quad "
        "\\text{where } w = \\text{softmax}(\\theta_w)$",
        italic=True, size=11)
    add_para(doc,
        "where $\\theta_w \\in \\mathbb{R}^4$ are learnable parameters initialized uniformly. In "
        "parallel, a lightweight MLP with two hidden layers refines the cue vector to capture "
        "nonlinear interactions between cues:")
    add_para(doc,
        "$F_i^{\\text{mlp}} = \\text{MLP}(\\mathbf{s}_i) = W_3 \\cdot \\text{GELU}(W_2 \\cdot "
        "\\text{GELU}(W_1 \\cdot \\mathbf{s}_i + b_1) + b_2) + b_3$",
        italic=True, size=11)
    add_para(doc,
        "The final foregroundness score combines both components and applies spatial smoothing:")
    add_para(doc,
        "$F_i = F_i^{\\text{linear}} + F_i^{\\text{mlp}} + 0.5 \\cdot "
        "\\sigma(W_s * \\text{Reshape}(F^{\\text{linear}} + F^{\\text{mlp}}))_{[i]}$",
        italic=True, size=11)
    add_para(doc,
        "where $W_s$ is a $3 \\times 3$ convolutional kernel initialized to uniform $1/9$ and "
        "$\\sigma$ denotes the sigmoid function. This three-component fusion strategy provides both "
        "interpretable linear contributions and flexible nonlinear refinement, with the spatial "
        "smoothing ensuring spatial coherence of the foregroundness map. The total ForegroundnessHead "
        "introduces fewer than 5K parameters.")

    add_sub(doc, "3.4", "Step 3: Adaptive Foreground Budget")
    add_para(doc,
        "A critical design decision is how to convert the continuous foregroundness scores "
        "$\\{F_i\\}_{i=1}^N$ into binary or soft foreground masks. Existing methods typically use "
        "fixed-ratio Top-K selection, where the top K% of patches are selected as foreground "
        "regardless of the image content. This rigid budget is problematic: images with a single "
        "small object may have fewer than 10% foreground patches, while images with multiple large "
        "objects may have over 50%.")
    add_para(doc,
        "We address this with an Adaptive Budget Module that predicts an image-specific threshold "
        "$\\tau$ from the global feature $c_{\\text{cls}}$, and applies a temperature-controlled "
        "sigmoid to produce a soft foreground mask:")
    add_para(doc,
        "$\\tau = \\text{MLP}_\\tau(c_{\\text{cls}}), \\quad "
        "m_i^F = \\sigma\\left(\\frac{F_i - \\tau}{T}\\right)$",
        italic=True, size=11)
    add_para(doc,
        "where $\\text{MLP}_\\tau$ is a two-layer MLP that maps the global feature to a scalar "
        "threshold, and $T$ is a learnable temperature parameter initialized to $0.1$ (in log-space "
        "for stable optimization) and clamped to $[0.01, 1.0]$ during forward passes. The temperature "
        "controls the sharpness of the mask: lower $T$ produces more binary-like masks, while higher $T$ "
        "yields softer, more gradual transitions between foreground and background.")
    add_para(doc,
        "The adaptive threshold mechanism has several advantages over fixed Top-K: (i) it naturally "
        "adapts to images with different foreground proportions; (ii) the soft mask avoids the "
        "discontinuity of hard thresholding, which creates artifacts at mask boundaries and destabilizes "
        "training gradients; and (iii) the learnable temperature enables the model to control mask "
        "sharpness as a function of training progress. During training, we apply a budget "
        "regularization loss (Section 3.8) to prevent degenerate solutions where the mask collapses "
        "to all-foreground or all-background.")

    add_sub(doc, "3.5", "Step 4: Dual CLS Decoupling")

    add_subsub(doc, "3.5.1  Foreground CLS Aggregation")
    add_para(doc,
        "Given the soft foreground mask $\\mathbf{m}^F = [m_1^F, ..., m_N^F]$, we compute the "
        "foreground CLS representation at each extracted layer by weighted averaging of patch features:")
    add_para(doc,
        "$B_l^F = \\frac{\\sum_{i=1}^{N} m_i^F \\cdot P_i^l}{\\sum_{i=1}^{N} m_i^F + \\epsilon}, "
        "\\quad B_l^F \\in \\mathbb{R}^{D}$",
        italic=True, size=11)
    add_para(doc,
        "This produces a sequence of foreground CLS tokens $\\mathbf{B}^F = [B_{l_1}^F, ..., B_{l_K}^F] "
        "\\in \\mathbb{R}^{K \\times D}$ that represent the foreground content as seen from each "
        "extracted layer. Crucially, these foreground CLS tokens maintain the multi-layer diversity of "
        "the backbone representations while focusing exclusively on foreground-relevant information.")

    add_subsub(doc, "3.5.2  Background CLS Aggregation")
    add_para(doc,
        "Rather than simply using $\\mathbf{m}^B = 1 - \\mathbf{m}^F$ as the background mask—which "
        "would make the background stream mechanically dependent on the foreground stream—we introduce "
        "a dedicated BackgroundnessHead that independently predicts a background mask from the final-layer "
        "patch features. This head consists of a two-layer MLP that outputs a background score, combined "
        "with an uncertainty predictor that estimates the reliability of each patch's backgroundness "
        "prediction:")
    add_para(doc,
        "$s_i^B = \\sigma(\\text{MLP}_B(P_i^{l_K})), \\quad "
        "u_i = \\sigma(\\text{MLP}_U(P_i^{l_K})), \\quad "
        "m_i^B = s_i^B \\cdot (1 - m_i^F) \\cdot u_i$",
        italic=True, size=11)
    add_para(doc,
        "The multiplicative composition ensures that: (i) the background mask excludes high-confidence "
        "foreground patches via the $(1 - m_i^F)$ term; (ii) the uncertainty term $u_i$ prevents "
        "ambiguous boundary patches from being confidently assigned to either stream; and (iii) the "
        "independent score $s_i^B$ allows the background stream to develop its own semantic preferences "
        "rather than being a passive complement of the foreground. The background CLS tokens "
        "$\\mathbf{B}^B = [B_{l_1}^B, ..., B_{l_K}^B]$ are then computed analogously:")
    add_para(doc,
        "$B_l^B = \\frac{\\sum_{i=1}^{N} m_i^B \\cdot P_i^l}{\\sum_{i=1}^{N} m_i^B + \\epsilon}$",
        italic=True, size=11)

    add_subsub(doc, "3.5.3  Why Dual Decoupling Rather Than Background Suppression")
    add_para(doc,
        "A natural question is why we maintain a separate background stream rather than simply "
        "suppressing background information. Three reasons motivate this design. First, background "
        "context provides complementary semantic information: scene type, spatial layout, and "
        "co-occurrence statistics that are valuable for classification and may even aid localization. "
        "Second, completely suppressing background forces all contextual information through the "
        "foreground stream, which may dilute the foreground signal—analogous to the information "
        "dilution problem in standard residual connections. Third, maintaining separate streams enables "
        "the task-adaptive fusion (Section 3.7) to learn how much background context each task needs, "
        "which varies substantially: object discovery benefits from minimal background, while "
        "classification may benefit from scene context.")

    add_sub(doc, "3.6", "Step 5: Depth-Wise Attention")
    add_para(doc,
        "After obtaining the dual CLS sequences $\\mathbf{B}^F$ and $\\mathbf{B}^B$, we perform "
        "independent depth-wise attention over the layer dimension for each stream. This mechanism "
        "directly adapts the Attention Residuals concept from language models to the visual domain, "
        "enabling each stream to selectively weight the contributions of different layers.")
    add_para(doc,
        "For the foreground stream, we compute:")
    add_para(doc,
        "$\\beta_l^F = \\frac{\\exp(w_F^T \\cdot \\text{RMSNorm}(B_l^F))}{\\sum_{j=1}^{K} "
        "\\exp(w_F^T \\cdot \\text{RMSNorm}(B_j^F))}, \\quad "
        "C_F = \\sum_{l=1}^{K} \\beta_l^F \\cdot B_l^F$",
        italic=True, size=11)
    add_para(doc,
        "where $w_F \\in \\mathbb{R}^D$ is a learned pseudo-query vector initialized to zero, "
        "ensuring uniform initial attention weights. RMSNorm is applied to the keys (block features) "
        "to prevent layers with large activation magnitudes from dominating the attention weights—a "
        "key design choice validated by the Attention Residuals ablation study. The background stream "
        "uses an analogous formulation with its own pseudo-query $w_B$, producing $C_B$ and "
        "$\\beta^B$.")
    add_para(doc,
        "The zero initialization of pseudo-queries is crucial: at the start of training, all layers "
        "contribute equally ($\\beta_l = 1/K$), so the initial behavior matches standard average "
        "pooling across layers. As training progresses, the model learns to emphasize layers that "
        "are most informative for each stream—typically later layers for semantic understanding and "
        "earlier layers for spatial precision.")

    add_sub(doc, "3.7", "Step 6: Task-Adaptive Fusion")
    add_para(doc,
        "The final image representation is produced by adaptively fusing the foreground and background "
        "streams. We introduce a learnable gate $g \\in [0, 1]$ that controls the relative contribution "
        "of each stream:")
    add_para(doc,
        "$g = \\sigma(\\text{MLP}_g([C_F; C_B; c_{\\text{cls}}])), \\quad "
        "C = g \\cdot C_F + (1 - g) \\cdot C_B$",
        italic=True, size=11)
    add_para(doc,
        "where $[\\cdot; \\cdot; \\cdot]$ denotes concatenation and $\\text{MLP}_g$ is a two-layer MLP "
        "mapping from $3D$ to 1. The gate takes all three representations as input—the foreground CLS "
        "$C_F$, the background CLS $C_B$, and the global backbone feature $c_{\\text{cls}}$—allowing it "
        "to make an informed decision about how much foreground vs. background information the task "
        "requires. For example, object discovery tasks (where objects are the primary target) may learn "
        "$g \\to 1$, while classification tasks (where scene context matters) may learn a more balanced "
        "gate value.")
    add_para(doc,
        "The fused representation $C$ is then passed to the appropriate task head:")
    add_para(doc,
        "Classification: A two-layer MLP maps $C$ to class logits. "
        "Segmentation: Patch features from the final layer are projected to the text embedding space "
        "(for OpenCLIP) or the CLS projection space, producing per-patch similarity scores that are "
        "reshaped to spatial maps. A learnable background bias based on $(1 - m^F)$ is added to the "
        "background class logits. "
        "Object Discovery: A scoring MLP produces per-patch scores, which are combined with the "
        "foreground mask and cosine similarity to $C$ for final score map prediction.")

    add_sub(doc, "3.8", "Training Objective")
    add_para(doc,
        "The total training loss combines a task-specific primary loss with three auxiliary losses "
        "that regularize the aggregation module:")
    add_para(doc,
        "$\\mathcal{L} = \\mathcal{L}_{\\text{task}} + \\lambda_{\\text{fg}} \\mathcal{L}_{\\text{mask}} "
        "+ \\lambda_{\\text{dec}} \\mathcal{L}_{\\text{decouple}} + \\lambda_{\\text{bud}} "
        "\\mathcal{L}_{\\text{budget}}$",
        italic=True, size=11)

    add_subsub(doc, "3.8.1  Task Loss")
    add_para(doc,
        "For object discovery, we supervise the foreground mask using pseudo-ground-truth masks derived "
        "from the initial foregroundness scores. Patches with scores exceeding one standard deviation "
        "above the mean are labeled as foreground, generating binary pseudo-masks that provide the "
        "training signal for the aggregation module. The task loss combines binary cross-entropy and "
        "Dice loss for robust segmentation supervision:")
    add_para(doc,
        "$\\mathcal{L}_{\\text{mask}} = \\text{BCE}(m^F, \\hat{m}) + (1 - \\text{Dice}(m^F, \\hat{m}))$",
        italic=True, size=11)
    add_para(doc,
        "For segmentation, standard cross-entropy loss between predicted segmentation logits and "
        "ground-truth masks is used. For classification, cross-entropy on class logits applies.")

    add_subsub(doc, "3.8.2  Decoupling Loss")
    add_para(doc,
        "To encourage the foreground and background streams to capture complementary rather than "
        "redundant information, we minimize the squared cosine similarity between $C_F$ and $C_B$:")
    add_para(doc,
        "$\\mathcal{L}_{\\text{decouple}} = \\left(\\frac{C_F \\cdot C_B}{\\|C_F\\| \\cdot "
        "\\|C_B\\|}\\right)^2$",
        italic=True, size=11)
    add_para(doc,
        "This loss pushes the two streams toward orthogonality in the feature space, ensuring that "
        "each stream encodes distinct semantic information. We use the squared formulation rather than "
        "the raw cosine similarity because the squared loss provides stronger gradients when the "
        "streams are nearly orthogonal (near zero), preventing the loss from becoming ineffective "
        "late in training. The loss weight is $\\lambda_{\\text{dec}} = 0.05$.")

    add_subsub(doc, "3.8.3  Budget Regularization")
    add_para(doc,
        "Without regularization, the adaptive budget may converge to degenerate solutions: all-"
        "foreground (where every patch receives high mask weight, negating the selection benefit) or "
        "all-background (where the foreground stream receives no information). We prevent these failure "
        "modes through a range penalty on the average foreground ratio $r = \\frac{1}{N}\\sum_i m_i^F$:")
    add_para(doc,
        "$\\mathcal{L}_{\\text{budget}} = \\max(0, r_{\\min} - r)^2 + \\max(0, r - r_{\\max})^2$",
        italic=True, size=11)
    add_para(doc,
        "where $r_{\\min} = 0.1$ and $r_{\\max} = 0.7$ are the lower and upper bounds on the "
        "acceptable foreground ratio. These bounds are chosen based on empirical analysis: typical "
        "object bounding boxes in VOC and COCO cover 10–60% of the image area, and allowing up to "
        "70% accommodates multi-object scenes while the 10% lower bound prevents mask collapse. "
        "The loss weight is $\\lambda_{\\text{bud}} = 0.01$.")

    add_sub(doc, "3.9", "Complexity Analysis")
    add_para(doc,
        "The DAFB-CLS aggregation module is designed to be lightweight and efficient. For a "
        "standard ViT-S/16 with $D = 384$, $N = 196$, and $K = 4$ extracted layers:")
    add_para(doc,
        "Feature extraction: The multi-layer hook mechanism adds negligible overhead (only "
        "intermediate tensor references), requiring no additional forward passes through the backbone. "
        "Foregroundness scoring: Four cue computations are each O(ND) or O(N), with the MLP fusion "
        "adding $O(N \\cdot H)$ where $H = 256$ is the hidden dimension. "
        "Adaptive budget: The threshold predictor MLP is $O(D \\cdot H + H)$, shared across all "
        "patches. "
        "Dual CLS aggregation: Two weighted averages over K layers, each $O(KND)$. "
        "Depth-wise attention: Two softmax attention computations with pseudo-queries, each $O(KD)$. "
        "Fusion: One MLP with input dimension $3D$, output dimension 1.")
    add_para(doc,
        "Total trainable parameters: 0.59M (DINO ViT-S/16) or 1.74M (OpenCLIP ViT-B/16), "
        "representing 2.7% and 2.0% of the respective frozen backbones. Inference latency "
        "overhead: 1.82ms (DINO) and 3.45ms (OpenCLIP), both under 31% relative overhead. "
        "Peak GPU memory scales linearly with the number of extracted layers due to storing "
        "multi-layer features for aggregation, resulting in approximately 2.8× (DINO) and "
        "3.0× (OpenCLIP) overhead relative to the baseline ViT forward pass.")


# ═══════════════════════════════════════════════════════════════
#   MAIN
# ═══════════════════════════════════════════════════════════════

def main():
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Times New Roman"
        hs.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        hs.font.color.rgb = RGBColor(0, 0, 0)

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    add_title(doc, "DAFB-CLS: Depth-Adaptive Foreground-Background CLS Decoupling for Vision Transformers", level=0)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(doc, "Related Work and Method Sections (Draft)", size=12, indent=0,
             space_before=6, space_after=24)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    write_related_work(doc)
    doc.add_page_break()
    write_methods(doc)

    out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                       "DAFB_CLS_RelatedWork_Methods.docx")
    doc.save(out)
    print(f"Saved to: {out}")


if __name__ == "__main__":
    main()
