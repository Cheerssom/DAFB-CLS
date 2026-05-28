"""
生成论文 Related Work 和 Method 章节（中文版）
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Microsoft YaHei'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    hs.font.color.rgb = RGBColor(0, 51, 102)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
    return table

def bp(doc, text):
    doc.add_paragraph(text, style='List Bullet')

def add_formula(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(80, 80, 80)
    return p

# ==================== RELATED WORK ====================
title = doc.add_heading('Related Work', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# --- 2.1 Vision Transformers ---
doc.add_heading('2.1 Vision Transformers', level=1)

doc.add_paragraph(
    'Vision Transformer（ViT）[Dosovitskiy等, 2021]首次将Transformer架构引入计算机视觉领域，'
    '将图像分割为固定大小的patch序列，通过自注意力机制建模patch之间的全局依赖关系。'
    'ViT在大规模预训练下展现出优于卷积网络的性能，催生了一系列后续工作。'
    'DeiT [Touvron等, 2021]通过知识蒸馏和数据增强策略，使ViT在中等规模数据集上也能有效训练。'
    '在自监督学习方面，DINO [Caron等, 2021]和DINOv2 [Oquab等, 2024]通过自蒸馏框架学习到了高质量的视觉表示，'
    '其自注意力图自然地包含了语义分割信息。在视觉-语言预训练方面，'
    'CLIP [Radford等, 2021]和OpenCLIP [Cherti等, 2023]通过对比学习将图像和文本映射到共享的嵌入空间，'
    '支持零样本分类和开放词汇识别。本文基于DINO和OpenCLIP两种预训练骨干，'
    '提出一种后处理聚合框架来改善其CLS token的空间定位能力。'
)

# --- 2.2 ViT中的注意力机制与可解释性 ---
doc.add_heading('2.2 ViT中的注意力机制与可解释性', level=1)

doc.add_paragraph(
    'ViT中的自注意力机制是其可解释性的核心。CLS token通过全局自注意力与所有patch token交互，'
    '最终汇聚为用于分类的全局表示。研究发现，CLS token的注意力图可以反映模型对图像不同区域的关注程度，'
    '因此被广泛用于可视化和弱监督定位。GradCAM [Selvaraju等, 2017]通过反向传播梯度加权激活图来生成类激活热力图，'
    'CAM [Zhou等, 2016]则利用全局平均池化层的权重直接生成激活图。'
    'DINO-seg [Caron等, 2021]发现DINO模型的自注意力图在最后一层自然地分割出前景目标，'
    '可以通过简单的阈值化获得语义分割掩码。TokenCut [Wang等, 2023]进一步将自注意力图建模为图结构，'
    '通过谱聚类实现无监督物体分割。然而，这些方法都依赖单一的注意力信号，'
    '缺乏对前景和背景信息的显式解耦，且未考虑深度维度的选择性聚合。'
)

# --- 2.3 Register Token与Artifact消除 ---
doc.add_heading('2.3 Register Token与Artifact消除', level=1)

doc.add_paragraph(
    'Darcet等[2024]首次系统性地揭示了DINOv2训练过程中出现的artifact现象：'
    '某些patch token会在训练过程中被推入特征空间的极端区域（高范数、低信息量），'
    '在注意力图中形成异常模式。作者将artifact分为两种类型：'
    'Type-1（单一patch接收过多注意力权重）和Type-2（注意力在所有token上均匀分布）。'
    '为解决这一问题，作者提出在训练时向输入序列添加额外的register token，'
    '让artifact注意力被这些专用容器吸收，推理时移除register token即可恢复清晰的注意力图。'
    '改进版方案进一步将artifact register替换为显式的前景和背景register，'
    '使token不仅吸收无用注意力，还承担语义区分功能。'
)

doc.add_paragraph(
    '然而，Register方法存在几个局限性：（1）需要从头训练或微调模型，无法即插即用到已有的预训练模型；'
    '（2）register token是无差别的占位符，没有显式的前景/背景语义区分能力；'
    '（3）所有层共享相同的register token，未考虑深度维度的选择性。'
    '本文的DAFB-CLS框架通过后处理方式，在推理阶段直接对冻结的预训练模型进行前景/背景解耦，'
    '无需重新训练backbone，同时引入深度自适应注意力机制。'
)

# --- 2.4 残差连接与深度聚合 ---
doc.add_heading('2.4 残差连接与深度聚合', level=1)

doc.add_paragraph(
    '标准Transformer的残差连接形式为 x_{l+1} = x_l + F_l(x_l)，'
    '即每一层的输出与输入进行等权重累加。当网络有L层时，最终输出是所有历史层输出的等权重叠加，'
    '这意味着模型无法自适应地决定哪些层的输出对当前任务最重要。'
    'Attention Residuals论文指出这一"惰性累加"（Lazy Accumulation）问题，'
    '并提出用注意力机制替代固定累加：对深度维度施加softmax注意力，'
    '让模型学习一个权重分布，自适应地聚合各层的贡献。'
)

doc.add_paragraph(
    '然而，Attention Residuals的方法仍是对单一CLS token的操作，'
    '没有将前景和背景信息分别处理，也缺乏显式的前景感知机制。'
    '本文在这一思想的基础上，为前景和背景流分别设计独立的深度注意力，'
    '并结合多线索前景区分和自适应掩码，实现了更精细的深度聚合。'
)

# --- 2.5 无监督物体发现与开放词汇分割 ---
doc.add_heading('2.5 无监督物体发现与开放词汇分割', level=1)

doc.add_paragraph(
    '无监督物体发现旨在在没有标注的情况下定位图像中的主要物体。'
    '传统方法依赖于显著性检测或超像素聚类。近年来，基于预训练ViT的方法取得了显著进展。'
    'LOST [Siméoni等, 2021]利用DINO的自注意力图中的特殊token（如[DINO] token）来定位物体。'
    'Tokencut [Wang等, 2023]将自注意力图建模为二部图，通过归一化割实现前景/背景分离。'
    'FreeSOLO [Wang等, 2022]通过自监督学习直接预测实例掩码。'
    'LaSt-ViT [Shi等, 2026]提出利用频率稳定性来区分前景和背景patch，'
    '通过FFT低通滤波计算每个patch的稳定性分数，选择最稳定的patch作为前景。'
    '然而，LaSt-ViT仅使用单一的频率信号和固定的Top-K选择，'
    '无法适应不同图像的前景面积差异，且硬阈值在掩码边界产生严重伪影。'
)

doc.add_paragraph(
    '开放词汇分割（Open-Vocabulary Segmentation）旨在分割任意文本描述的目标类别。'
    'MaskCLIP [Zhou等, 2022]利用CLIP的patch-text相似度生成类别无关的掩码提议。'
    'OVSeg [Liang等, 2023]通过微调CLIP的视觉编码器来改善开放词汇分割性能。'
    '本文的DAFB-CLS框架可以同时支持无监督物体发现和开放词汇分割两种任务，'
    '通过任务自适应融合门控来平衡前景和背景信息的贡献。'
)

# --- 2.6 后处理方法 ---
doc.add_heading('2.6 后处理方法（Post-hoc Methods）', level=1)

doc.add_paragraph(
    '后处理方法是指在不修改预训练模型参数的前提下，通过附加的轻量模块来改善模型的特定能力。'
    '这类方法的优势在于：（1）不改变预训练模型的任何参数，避免灾难性遗忘；'
    '（2）训练成本极低，因为附加模块的参数量远小于backbone；'
    '（3）可以即插即用到任何已有的预训练模型上。'
    '本文的DAFB-CLS是一种典型的后处理方法：backbone完全冻结，'
    '仅训练前景/背景聚合器、深度注意力和融合门控等轻量模块。'
    '这使得DAFB-CLS可以直接应用于DINO、OpenCLIP等社区中广泛使用的预训练模型，'
    '无需重新训练backbone。'
)

# ==================== METHOD ====================
doc.add_page_break()
title = doc.add_heading('Method', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# --- 3.1 概述 ---
doc.add_heading('3.1 总体框架概述', level=1)

doc.add_paragraph(
    'DAFB-CLS（Depth-Adaptive Foreground-Background CLS Decoupling）是一个后处理聚合框架，'
    '旨在改善Vision Transformer中CLS token的空间定位能力和可解释性。'
    '给定一个冻结的预训练ViT骨干（如DINO或OpenCLIP），'
    'DAFB-CLS将CLS token解耦为独立的前景语义流和背景语义流，'
    '并通过深度自适应注意力机制分别聚合各层的贡献，'
    '最终通过任务自适应门控融合生成最终的CLS表示。'
)

doc.add_paragraph(
    '整个框架由五个阶段组成：'
    '（1）多线索前景区分度计算——融合4种互补线索判断哪些patch属于前景；'
    '（2）自适应前景预算——将连续的前景性分数转换为自适应软掩码；'
    '（3）双CLS解耦聚合——分别计算前景和背景CLS token；'
    '（4）深度注意力——前景和背景流各自独立地进行深度维度的注意力聚合；'
    '（5）任务自适应融合——通过门控机制自适应地平衡前景和背景的贡献。'
)

doc.add_paragraph('形式化地，给定输入图像I，框架的前向传播过程为：')

add_formula(doc, '{x^l}_{l=1}^L, c = Backbone(I)                           -- 特征提取')
add_formula(doc, 'm^F, m^B = CueFusion(x^L) + AdaptiveBudget(...)          -- 前景/背景掩码')
add_formula(doc, 'B_l^F = Agg(x^l, m^F),  B_l^B = Agg(x^l, m^B)          -- 双CLS聚合')
add_formula(doc, 'C_F = DepthAttn_F({B_l^F}), C_B = DepthAttn_B({B_l^B})  -- 深度注意力')
add_formula(doc, 'C = Gate(C_F, C_B, c)                                     -- 任务自适应融合')

doc.add_paragraph(
    '其中x^l是第l层的patch特征，c是全局CLS token，L是提取的层数（默认4层：第3、6、9、12层）。'
    'm^F和m^B分别是前景和背景掩码，B_l^F和B_l^B是第l层的前景/背景CLS token，'
    'C_F和C_B是深度聚合后的前景/背景CLS汇总，C是最终输出。'
)

# --- 3.2 多线索前景区分 ---
doc.add_heading('3.2 多线索前景区分度计算', level=1)

doc.add_paragraph(
    '现有的前景区分方法（如LaSt-ViT）仅依赖单一的频率稳定性信号，'
    '这在平滑背景（天空、墙面）场景下会失效。'
    '我们设计了4种互补的前景区分线索，每种基于不同的归纳偏置，'
    '并通过可学习的融合机制将它们组合为统一的前景性分数。'
)

doc.add_heading('3.2.1 频率稳定性线索（FrequencyStabilityCue）', level=2)

doc.add_paragraph(
    '该线索的直觉是：前景目标的特征在频率空间中包含更多高频成分（边缘、纹理），'
    '而背景区域的特征更平滑，能量集中在低频。给定第l层第i个patch的特征向量x_i^l ∈ R^D，'
    '频率稳定性分数的计算过程如下：'
)

doc.add_paragraph('首先，对特征进行FFT变换，并应用高斯低通滤波器：')
add_formula(doc, 'x_fft = FFT(x_i^l)')
add_formula(doc, 'H(f) = exp(-0.5 * (f / sigma)^2)       -- sigma = 0.25')
add_formula(doc, 'x_filtered = IFFT(x_fft * H(f)).real')

doc.add_paragraph('然后，计算稳定性分数：')
add_formula(doc, 's_freq = mean_D( x_filtered / (|x_filtered - x_i^l| + eps) )')

doc.add_paragraph(
    '当patch特征平滑时（如背景），x_filtered ≈ x_i^l，分母很小，稳定性分数高；'
    '当patch特征包含高频成分时（如前景边缘），|x_filtered - x_i^l|大，稳定性分数低。'
    '我们取最后一层的稳定性分数作为该线索的输出。'
)

doc.add_paragraph(
    '可选的纹理互补模式：额外计算高频残差能量 texture_energy = ||x_i^l - x_filtered||_2，'
    '与稳定性分数通过可学习权重 w = sigmoid(w_param) 进行混合：'
    's_freq = w * texture_norm + (1-w) * stability。'
    '这有助于区分"平滑前景"（如白色汽车）和"平滑背景"（如天空）。'
)

doc.add_heading('3.2.2 深度一致性线索（DepthConsistencyCue）', level=2)

doc.add_paragraph(
    '该线索的核心假设是：前景目标的patch在不同层的特征表示应该保持较高的一致性，'
    '因为它们对应同一个语义实体；而背景patch的特征在不同层可能变化较大。'
    '给定patch i在所有L层的特征 {x_i^1, ..., x_i^L}，计算过程如下：'
)

add_formula(doc, 'x_mean = mean_L(x_i^l)')
add_formula(doc, 's_depth = mean_L( cosine(x_i^l, x_mean) )')

doc.add_paragraph(
    '其中cosine(·,·)是余弦相似度。前景patch的跨层一致性高，得分高；'
    '背景patch的跨层一致性低，得分低。'
)

doc.add_heading('3.2.3 语义对齐线索（SemanticAlignmentCue）', level=2)

doc.add_paragraph(
    '该线索衡量每个patch与CLS token的语义相关性。'
    'CLS token编码了图像的全局语义信息，与CLS token相似度高的patch更可能是前景。'
    '计算过程为：'
)

add_formula(doc, 's_sem = mean_L( cosine(x_i^l, c) )')

doc.add_paragraph(
    '其中c是全局CLS token。对于OpenCLIP骨干，还可使用patch-text相似度：'
    's_sem = max_k( cosine(x_i^L, t_k) )，其中t_k是第k个类别的文本嵌入。'
)

doc.add_heading('3.2.4 空间紧凑性线索（SpatialCompactnessCue）', level=2)

doc.add_paragraph(
    '该线索利用"前景目标通常在空间上是紧凑的"这一先验。'
    '将语义对齐分数reshape为14×14的2D图，通过3×3平均池化进行空间平滑：'
)

add_formula(doc, 's_spatial = AvgPool2d( reshape(s_sem, 14, 14), kernel=3 )')

doc.add_paragraph(
    '平滑操作抑制孤立的高分patch，鼓励空间连续的前景区域。'
)

doc.add_heading('3.2.5 线索融合', level=2)

doc.add_paragraph(
    '4条线索在ForegroundnessHead中通过双重机制融合：'
)

doc.add_paragraph('（1）线性融合：可学习的softmax归一化权重：')
add_formula(doc, 'w = softmax(w_1, w_2, w_3, w_4)')
add_formula(doc, 'F_linear = w_1 * s_freq + w_2 * s_depth + w_3 * s_sem + w_4 * s_spatial')

doc.add_paragraph('（2）非线性修正：将4条线索拼接为4维向量，通过3层MLP：')
add_formula(doc, 'F_mlp = MLP([s_freq; s_depth; s_sem; s_spatial])    -- [4->256->256->1]')

doc.add_paragraph('（3）空间平滑残差：3×3 Conv2d(1,1)，权重初始化为1/9：')
add_formula(doc, 'F_smooth = Sigmoid(Conv2d(reshape(F_combined, 14, 14)))')

doc.add_paragraph('最终前景性分数：')
add_formula(doc, 'F = F_linear + F_mlp + 0.5 * F_smooth')

# --- 3.3 自适应前景预算 ---
doc.add_heading('3.3 自适应前景预算', level=1)

doc.add_paragraph(
    '将连续的前景性分数F转换为二值化的前景掩码m^F，同时保持可微性。'
    '这是对LaSt-ViT固定Top-K选择的关键改进——固定比例无法适应不同图像的前景面积差异。'
)

doc.add_paragraph('自适应阈值预测：给定全局CLS特征c，通过MLP预测每张图像的阈值tau：')
add_formula(doc, 'tau = MLP_tau(c)                     -- Linear(D, 128) -> GELU -> Linear(128, 1)')

doc.add_paragraph('可学习温度参数T控制sigmoid的陡峭程度：')
add_formula(doc, 'T = exp(log_T).clamp(0.01, 1.0)     -- 初始值 T = 0.1')

doc.add_paragraph('软前景掩码：')
add_formula(doc, 'm^F = sigmoid((F - tau) / T)')

doc.add_paragraph(
    '当T很小时，掩码接近二值化；当T较大时，掩码更平滑。'
    'tau和T都是可学习的，使得掩码阈值能够自适应于每张图像的内容。'
)

doc.add_paragraph(
    '独立的BackgroundnessHead从最后一层patch特征直接预测背景掩码：'
)
add_formula(doc, 'bg_score = sigmoid(MLP_bg(x^L))')
add_formula(doc, 'uncertainty = sigmoid(MLP_unc(x^L))')
add_formula(doc, 'm^B = bg_score * (1 - m^F) * uncertainty')

doc.add_paragraph(
    '背景掩码是三个因子的乘积：（1）学习的背景概率，（2）前景掩码的补集（防止重叠），'
    '（3）不确定性门控（降低低置信度区域的影响）。'
)

# --- 3.4 双CLS解耦聚合 ---
doc.add_heading('3.4 双CLS解耦聚合', level=1)

doc.add_paragraph(
    '核心创新。不同于Register论文的被动消除背景和Attention Residuals的单一深度聚合，'
    '我们分别计算前景和背景CLS token。对每个提取的ViT层独立进行加权聚合：'
)

add_formula(doc, 'B_l^F = sum_i(m_i^F * x_i^l) / sum_i(m_i^F)    -- 前景CLS token')
add_formula(doc, 'B_l^B = sum_i(m_i^B * x_i^l) / sum_i(m_i^B)    -- 背景CLS token')

doc.add_paragraph(
    '聚合后的B^F和B^B各自是[B, L, D]的张量。前景CLS token只包含前景区域的信息，'
    '背景CLS token只包含背景区域的信息。这种显式解耦确保了后续的深度注意力'
    '不会受到背景信息的干扰。'
)

doc.add_paragraph(
    '消融实验表明，双CLS解耦对空间定位至关重要：'
    '去除双CLS后PiB从45.07%降至33.13%（-11.94pp），'
    '证明分别维护前景和背景token对空间定位有显著贡献。'
)

# --- 3.5 深度注意力 ---
doc.add_heading('3.5 深度自适应注意力', level=1)

doc.add_paragraph(
    '受Attention Residuals启发，我们为前景和背景分别设计独立的深度注意力。'
    '每个流维护一个可学习的query向量，对L层的CLS token进行softmax注意力加权：'
)

add_formula(doc, 'keys = RMSNorm(B^F)            -- [B, L, D]')
add_formula(doc, 'scores = q_F^T * keys          -- [B, L]')
add_formula(doc, 'beta_F = softmax(scores)       -- [B, L]')
add_formula(doc, 'C_F = sum(beta_F * B^F)        -- [B, D]')

doc.add_paragraph(
    '背景流类似，使用独立的query向量q_B和RMSNorm。'
    'query向量初始化为全零向量，使得初始时所有层的注意力权重相等（=1/L），'
    '模型行为退化为标准的等权平均——这是一种安全的起始点，'
    '训练过程中逐渐学习差异化的权重。'
)

doc.add_paragraph(
    'RMSNorm（Root Mean Square Normalization）用于归一化keys，公式为：'
)
add_formula(doc, 'RMSNorm(x) = x / sqrt(mean(x^2) + eps) * weight')

doc.add_paragraph(
    '相比LayerNorm，RMSNorm更简洁稳定，遵循现代LLM的设计实践。'
    '归一化在float32精度下进行以确保数值稳定。'
)

doc.add_paragraph(
    '消融实验表明，独立的前景/背景深度注意力优于共享注意力：'
    '共享后PiB从45.07%降至28.36%（-16.71pp），'
    '证明前景和背景信息需要不同的深度权重分布。'
)

# --- 3.6 任务自适应融合 ---
doc.add_heading('3.6 任务自适应融合', level=1)

doc.add_paragraph(
    '最终通过门控机制自适应地融合前景和背景CLS token：'
)

add_formula(doc, 'g = sigmoid(MLP([C_F; C_B; c]))       -- [3D -> 256 -> 1]')
add_formula(doc, 'C = g * C_F + (1 - g) * C_B')

doc.add_paragraph(
    '门控网络是一个3层MLP，输入是C_F、C_B和backbone的全局CLS token c'
    '三者拼接后的3D维向量。输出g是一个0-1之间的标量。'
    '当g趋近1时，最终输出以前景为主；当g趋近0时，以背景为主。'
    '这允许模型根据具体图像和任务自适应地调整前景/背景的贡献比例。'
)

# --- 3.7 任务头 ---
doc.add_heading('3.7 任务头', level=1)

doc.add_paragraph('DAFB-CLS支持三种下游任务，每种使用不同的任务头：')

doc.add_paragraph()
r = doc.add_paragraph().add_run('分类头（ClassificationHead）：')
r.bold = True
doc.add_paragraph(
    '两层MLP：Linear(D, 512) -> GELU -> Linear(512, num_classes)。'
    '输入为融合后的CLS token C，输出为类别logits。'
)

doc.add_paragraph()
r = doc.add_paragraph().add_run('分割头（SegmentationHead）：')
r.bold = True
doc.add_paragraph(
    '基于patch-text相似度的开放词汇分割。'
    '将patch特征通过线性投影映射到文本嵌入空间，'
    '计算patch与各类别文本嵌入的余弦相似度作为分割logits：'
)
add_formula(doc, 'patch_proj = normalize(Linear(D, text_dim)(x^L))')
add_formula(doc, 'seg_logits = patch_proj * text_features^T')
doc.add_paragraph(
    '对于背景类别，额外添加前景掩码的偏置：seg_logits[:, 0] += (1 - m^F) * seg_bg_scale，'
    '其中seg_bg_scale是可学习参数。最终logits乘以CLIP风格的温度缩放。'
)

doc.add_paragraph()
r = doc.add_paragraph().add_run('物体发现头（ObjectDiscoveryScoringHead）：')
r.bold = True
doc.add_paragraph(
    '结合学习的patch分数和CLS相似度，通过前景掩码加权：'
)
add_formula(doc, 'score = MLP(x^L) + cosine(x^L, c)')
add_formula(doc, 'score = score * m^F')

# --- 3.8 损失函数 ---
doc.add_heading('3.8 损失函数', level=1)

doc.add_paragraph('总训练损失由四项组成：')
add_formula(doc, 'L_total = L_task + lambda_fg * L_mask + lambda_decouple * L_decouple + lambda_budget * L_budget')

doc.add_paragraph()
r = doc.add_paragraph().add_run('主任务损失 L_task：')
r.bold = True
doc.add_paragraph(
    '分类任务使用交叉熵损失；分割任务使用patch-text相似度的交叉熵损失；'
    '物体发现任务使用patch得分与GT边界框的交叉熵损失。'
)

doc.add_paragraph()
r = doc.add_paragraph().add_run('前景掩码损失 L_mask：')
r.bold = True
doc.add_paragraph(
    'BCE损失与Dice损失的组合，监督预测的前景掩码与伪GT掩码的一致性。'
    '伪GT掩码从patch分数生成：选择分数高于均值+标准差的patch作为前景。'
)
add_formula(doc, 'L_mask = BCE(m^F, m_gt) + (1 - Dice(m^F, m_gt))')

doc.add_paragraph()
r = doc.add_paragraph().add_run('解耦损失 L_decouple：')
r.bold = True
doc.add_paragraph(
    '使用余弦平方损失推远前景和背景CLS token的方向：'
)
add_formula(doc, 'L_decouple = mean( cosine(C_F, C_B)^2 )')

doc.add_paragraph()
r = doc.add_paragraph().add_run('预算正则化损失 L_budget：')
r.bold = True
doc.add_paragraph(
    '将前景比例约束在[r_min, r_max]范围内，防止掩码坍缩：'
)
add_formula(doc, 'r = mean(m^F)')
add_formula(doc, 'L_budget = mean( max(r_min - r, 0)^2 + max(r - r_max, 0)^2 )')

doc.add_paragraph(
    '其中r_min = 0.1，r_max = 0.7。'
    '消融实验表明，去除预算正则化会导致Mask IoU从31.27%暴跌至13.91%。'
)

# --- 3.9 训练策略 ---
doc.add_heading('3.9 训练策略', level=1)

doc.add_paragraph(
    'DAFB-CLS采用后训练策略：backbone完全冻结，仅训练聚合器模块。'
    '可训练参数通过排除所有包含"extractor"的参数来实现。'
    '训练使用AdamW优化器，余弦退火学习率调度，梯度裁剪（max_norm=1.0）。'
)

doc.add_paragraph('对于不同骨干的超参数配置：')

add_table(doc,
    ['超参数', 'DINO ViT-S/16', 'OpenCLIP ViT-B/16'],
    [
        ['学习率', '1e-4 (VOC) / 5e-5 (COCO)', '5e-5'],
        ['batch_size', '16', '16'],
        ['epochs', '50', '50'],
        ['lambda_fg', '1.0 (VOC) / 0.05 (COCO)', '2.0'],
        ['lambda_decouple', '0.05', '0.1'],
        ['lambda_budget', '0.01', '0.02'],
        ['AMP', '关闭（COCO）/ 开启（VOC）', '开启'],
    ]
)

doc.add_paragraph(
    '对于COCO数据集，由于标注更复杂（80类、多实例），'
    '需要降低lambda_fg（0.05 vs 1.0）和学习率（5e-5 vs 1e-4），'
    '并关闭AMP以避免fp16精度溢出导致的训练崩溃。'
)

# Save
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'DAFB_CLS_论文章节_RelatedWork_Method.docx')
doc.save(output_path)
print(f'已保存: {output_path}')
